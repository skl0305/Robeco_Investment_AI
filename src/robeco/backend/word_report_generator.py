#!/usr/bin/env python3
"""
Robeco HTML-to-Word Report Generator
Converts generated HTML reports to Word documents (.docx) with exact layout preservation
Maintains Robeco professional styling, metrics grids, and institutional formatting
"""

import logging
import re
import io
import base64
from datetime import datetime
from typing import Dict, List, Optional, Any, Tuple
from pathlib import Path

# Word document generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION_START
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# HTML parsing
from bs4 import BeautifulSoup, NavigableString
import requests
from io import BytesIO
from PIL import Image

# Import InlineShape for image insertion
try:
    from docx.shared import Mm
    from docx.enum.dml import MSO_THEME_COLOR_INDEX
except ImportError:
    pass  # Optional imports

logger = logging.getLogger(__name__)

class RobecoWordReportGenerator:
    """
    Professional HTML-to-Word conversion for Robeco investment reports
    Preserves exact layout, styling, and structure from HTML reports
    """
    
    def __init__(self):
        # Image cache for company logos
        self._image_cache = {}
        
        # Flags for one-time elements
        self._first_analysis_item = False
        self._robeco_logo_added = False
        
        # Professional font configuration
        self.primary_font = 'Calibri'  # Fallback for Taz Semilight
        
        # Robeco brand colors (RGB values)
        self.robeco_colors = {
            'blue': RGBColor(0, 95, 144),         # #005F90
            'blue_darker': RGBColor(0, 61, 90),    # #003D5A
            'brown_black': RGBColor(59, 49, 42),   # #3B312A
            'orange': RGBColor(255, 140, 0),       # #FF8C00
            'text_dark': RGBColor(0, 0, 0),        # #000000
            'text_secondary': RGBColor(51, 51, 51), # #333333
            'accent_green': RGBColor(46, 125, 50),  # #2E7D32
            'accent_red': RGBColor(198, 40, 40),    # #C62828
        }
        
        # Font sizes mapping from HTML to Word (pt to pt conversion)
        self.font_size_mapping = {
            '57px': 57,     # report-title
            '27px': 27,     # report-subtitle  
            '25.5px': 26,   # section-title
            '20px': 20,     # bold-black-header.report-prose
            '18px': 18,     # bold-black-header, body
            '16.5pt': 17,   # report-footer
            '10pt': 10,     # metrics labels
        }
        
        logger.info("🏗️ Robeco Word Report Generator initialized")
    
    async def convert_html_to_word(
        self, 
        html_content: str, 
        company_name: str, 
        ticker: str,
        output_path: Optional[str] = None
    ) -> str:
        """
        Convert HTML report to Word document with exact layout preservation
        
        Args:
            html_content: Complete HTML report content
            company_name: Company name for document properties
            ticker: Stock ticker symbol
            output_path: Optional custom output path
            
        Returns:
            str: Path to generated Word document
        """
        logger.info(f"🔄 Converting HTML report to Word: {ticker}")
        
        try:
            # Parse HTML content
            soup = BeautifulSoup(html_content, 'lxml')
            
            # Create new Word document
            doc = Document()
            
            # Set document properties
            self._set_document_properties(doc, company_name, ticker)
            
            # Configure page layout
            self._configure_page_layout(doc)
            
            # Process presentation container or find slides directly
            presentation = soup.find('div', class_='presentation-container')
            if presentation:
                # ENHANCED SLIDE DETECTION: Find ALL slide variations used in 3-call architecture
                slides = []
                
                # Method 1: Find all divs with "slide" in their class attribute
                all_slide_divs = presentation.find_all('div', class_=lambda c: c and 'slide' in ' '.join(c))
                slides.extend(all_slide_divs)
                
                # Method 1.5: SPECIFIC ID SEARCH - Find slides by expected IDs (critical for 3-call architecture)
                expected_slide_ids = ['portrait-page-1', 'portrait-page-1A', 'investment-highlights-pitchbook-page', 
                                     'catalyst-page', 'company-analysis-page', 'financial-highlights-table-page',
                                     'slide-industry-analysis-part1', 'slide-financial-income-statement', 
                                     'slide-financial-balance-sheet', 'cash-flow-page', 'dcf-analysis-page', 
                                     'bull-bear-analysis-comprehensive']
                
                for slide_id in expected_slide_ids:
                    # First try within presentation container
                    slide_by_id = presentation.find('div', id=slide_id)
                    if slide_by_id:
                        slides.append(slide_by_id)
                        logger.info(f"🎯 DIRECT ID MATCH (in presentation): Found slide with ID '{slide_id}'")
                    else:
                        # If not found in presentation, search entire document
                        slide_by_id = soup.find('div', id=slide_id)
                        if slide_by_id:
                            slides.append(slide_by_id)
                            logger.info(f"🎯 DIRECT ID MATCH (in document): Found slide with ID '{slide_id}'")
                
                # Method 2: Backup CSS selectors for specific patterns
                backup_slides = (
                    presentation.select('div.slide') +              # Basic slide
                    presentation.select('div.slide.report-prose') + # Slide with report-prose  
                    presentation.select('div[class*="slide"]') +    # Any div containing "slide"
                    presentation.select('div[id*="page"]')          # Slides with page IDs
                )
                slides.extend(backup_slides)
                
                logger.info(f"🔍 ENHANCED DETECTION: Found {len(all_slide_divs)} slide divs, {len(backup_slides)} backup slides")
                
            else:
                # If no presentation container, find slides directly in body
                slides = []
                
                # Method 1: Find all divs with "slide" in their class attribute
                all_slide_divs = soup.find_all('div', class_=lambda c: c and 'slide' in ' '.join(c))
                slides.extend(all_slide_divs)
                
                # Method 2: Backup selectors
                backup_slides = (
                    soup.select('div.slide') + 
                    soup.select('div.slide.report-prose') + 
                    soup.select('div[class*="slide"]') +
                    soup.select('div[id*="page"]')
                )
                slides.extend(backup_slides)
                
                logger.info(f"🔍 ENHANCED DETECTION (no container): Found {len(all_slide_divs)} slide divs, {len(backup_slides)} backup slides")
            
            # DEBUG: Show raw slides before deduplication
            logger.info(f"🔍 RAW SLIDES BEFORE DEDUP: {len(slides)} total")
            raw_slide_info = []
            for i, slide in enumerate(slides):
                slide_id = slide.get('id', f'no-id-{i}')
                slide_classes = slide.get('class', [])
                raw_slide_info.append(f"{slide_id}({' '.join(slide_classes)})")
            logger.info(f"   📋 Raw slide details: {raw_slide_info}")
            
            # Remove duplicates while preserving order using element IDs and content
            seen_elements = set()
            unique_slides = []
            for slide in slides:
                # Create unique identifier using element ID or content hash
                slide_id = slide.get('id') or str(hash(str(slide)[:100]))
                if slide_id not in seen_elements:
                    seen_elements.add(slide_id)
                    unique_slides.append(slide)
                    logger.info(f"✅ KEPT SLIDE: {slide_id} ({slide.get('class', [])})")
                else:
                    logger.info(f"🔄 SKIPPED DUPLICATE: {slide_id} ({slide.get('class', [])})")
            
            slides = unique_slides
            
            logger.info(f"📄 Processing {len(slides)} slides")
            
            # ENHANCED DEBUG: Log comprehensive slide information
            slide_info = []
            for i, slide in enumerate(slides):
                slide_id = slide.get('id', f'no-id-{i}')
                slide_classes = slide.get('class', [])
                slide_info.append(f"{slide_id}({' '.join(slide_classes)})")
            
            logger.info(f"🔍 ENHANCED SLIDE DETECTION RESULTS:")
            logger.info(f"   📄 Total slides found: {len(slides)}")
            logger.info(f"   📋 Slide details: {slide_info}")
            
            # Debug: Log HTML structure summary if we found fewer slides than expected
            expected_min_slides = 10  # We expect at least 10 slides from 3-call architecture
            if len(slides) < expected_min_slides:
                logger.warning(f"⚠️ Only found {len(slides)} slides (expected >= {expected_min_slides}), analyzing HTML...")
                
                # Count all divs with any slide-related content
                all_divs = soup.find_all('div')
                slide_related_divs = [div for div in all_divs if div.get('class') and any('slide' in cls for cls in div.get('class', []))]
                page_id_divs = [div for div in all_divs if div.get('id') and 'page' in div.get('id', '')]
                portrait_divs = [div for div in all_divs if div.get('id') and 'portrait' in div.get('id', '')]
                
                logger.info(f"🔍 HTML ANALYSIS:")
                logger.info(f"   📄 Total divs in HTML: {len(all_divs)}")
                logger.info(f"   📄 Slide-related divs: {len(slide_related_divs)}")
                logger.info(f"   📄 Page ID divs: {len(page_id_divs)}")
                logger.info(f"   📄 Portrait divs: {len(portrait_divs)}")
                
                # Log the specific IDs found
                portrait_ids = [div.get('id') for div in portrait_divs]
                logger.info(f"   📋 Portrait IDs found: {portrait_ids}")
                logger.info(f"   📄 Divs with 'slide' in class: {len(slide_related_divs)}")
                logger.info(f"   📄 Divs with 'page' in ID: {len(page_id_divs)}")
                logger.info(f"   📄 Total HTML length: {len(html_content):,} characters")
                
                # Show first few div classes for debugging
                div_classes = [div.get('class', []) for div in all_divs[:15]]
                logger.info(f"   📋 First 15 div classes: {div_classes}")
                
                # Look for specific slide IDs we expect
                expected_ids = ['portrait-page-1', 'portrait-page-1A', 'investment-highlights-pitchbook-page', 
                               'catalyst-page', 'company-analysis-page', 'financial-highlights-table-page']
                found_ids = []
                for expected_id in expected_ids:
                    if soup.find('div', id=expected_id):
                        found_ids.append(expected_id)
                
                logger.info(f"   ✅ Expected slide IDs found: {found_ids}")
                logger.info(f"   ❌ Missing slide IDs: {set(expected_ids) - set(found_ids)}")
                
                # Check for presentation container
                if presentation:
                    logger.info("   ✅ Found presentation-container")
                    pres_divs = presentation.find_all('div')
                    logger.info(f"   📄 Divs inside presentation-container: {len(pres_divs)}")
                else:
                    logger.warning("   ⚠️ No presentation-container found, searching entire document")
            
            # Initialize header flag for the document
            self._header_added = False
            
            for i, slide in enumerate(slides):
                logger.info(f"🎯 Processing slide {i+1}/{len(slides)}")
                
                if i > 0:
                    # Add page break between slides
                    self._add_page_break(doc)
                
                # Process slide content
                self._process_slide(doc, slide, i+1)
            
            # Generate output filename
            if not output_path:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_path = f"/tmp/{ticker}_Investment_Report_{timestamp}.docx"
            
            # Save document
            doc.save(output_path)
            logger.info(f"✅ Word document generated: {output_path}")
            
            return output_path
            
        except Exception as e:
            logger.error(f"❌ HTML to Word conversion failed: {e}")
            raise
    
    def _set_document_properties(self, doc: Document, company_name: str, ticker: str):
        """Set document properties and metadata"""
        core_props = doc.core_properties
        core_props.title = f"{company_name} ({ticker}) - Investment Analysis"
        core_props.author = "Robeco Investment Analysis Platform"
        core_props.subject = f"Professional Investment Analysis - {ticker}"
        core_props.created = datetime.now()
        
    def _configure_page_layout(self, doc: Document):
        """Configure page layout to match HTML dimensions"""
        # Get page dimensions from HTML (1620px width, 2291px height)
        # Convert to inches (assuming 96 DPI: 1620px = 16.875", 2291px = 23.86")
        section = doc.sections[0]
        section.page_width = Inches(16.875)
        section.page_height = Inches(23.86)
        section.top_margin = Inches(1.09)    # 105px padding
        section.bottom_margin = Inches(0.47) # 45px footer space
        section.left_margin = Inches(1.02)   # 98px padding
        section.right_margin = Inches(1.02)  # 98px padding
        
    def _add_page_break(self, doc: Document):
        """Add page break between slides"""
        from docx.enum.text import WD_BREAK
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_break(WD_BREAK.PAGE)
        
    def _process_slide(self, doc: Document, slide_soup, slide_number: int):
        """Process individual slide content following exact HTML structure"""
        logger.info(f"🎯 Processing slide {slide_number}")
        
        # Debug: Log slide structure
        slide_classes = slide_soup.get('class', [])
        slide_id = slide_soup.get('id', 'no-id')
        logger.info(f"🔍 Slide classes: {slide_classes}, ID: {slide_id}")
        
        # STEP 1: Process slide-logo FIRST (HTML: first element in slide)
        slide_logo = slide_soup.find('div', class_='slide-logo')
        if slide_logo:
            logger.info("🏢 Processing slide-logo (top-right positioning)")
            self._add_slide_logo_header(doc, slide_logo)
        
        # Initialize header flag for this instance
        if not hasattr(self, '_header_added'):
            self._header_added = False
        
        # Process slide header if exists
        header = slide_soup.find('div', class_='slide-header')
        if header:
            logger.info("✅ Found slide-header")
            self._process_slide_header(doc, header)
        else:
            logger.info("ℹ️ No slide-header found")
        
        # Process slide content if exists
        content = slide_soup.find('div', class_='slide-content')
        if content:
            logger.info("✅ Found slide-content")
            self._process_slide_content(doc, content, slide_classes)
        else:
            logger.info("ℹ️ No slide-content found, processing slide directly")
            # If no slide-content wrapper, process the slide itself
            self._process_slide_content(doc, slide_soup, slide_classes)
        
        # Process report footer if exists
        footer = slide_soup.find('div', class_='report-footer')
        if footer:
            logger.info("✅ Found report-footer")
            self._process_report_footer(doc, footer)
        else:
            logger.info("ℹ️ No report-footer found")
    
    def _process_slide_header(self, doc: Document, header_soup):
        """Process slide header with title and subtitle"""
        # Report title
        title = header_soup.find(class_='report-title')
        if title:
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = p.add_run(title.get_text().strip())
            run.font.size = Pt(57)
            run.font.name = self.primary_font
            run.font.bold = True
            run.font.color.rgb = self.robeco_colors['brown_black']
            p.space_after = Pt(0)
        
        # Report subtitle
        subtitle = header_soup.find(class_='report-subtitle')
        if subtitle:
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = p.add_run(subtitle.get_text().strip())
            run.font.size = Pt(27)
            run.font.name = self.primary_font
            run.font.color.rgb = self.robeco_colors['text_dark']
            p.space_after = Pt(22)
    
    def _process_slide_content(self, doc: Document, content_soup, slide_classes=None):
        """Process main slide content"""
        logger.info(f"🔍 Processing slide content, type: {type(content_soup)}")
        
        # Initialize slide_classes if not provided
        if slide_classes is None:
            slide_classes = []
        
        # Check if content_soup has children (BeautifulSoup element) or needs different processing
        if hasattr(content_soup, 'children'):
            elements_processed = 0
            for element in content_soup.children:
                if hasattr(element, 'name') and element.name is not None:
                    elements_processed += 1
                    element_classes = element.get('class', [])
                    logger.info(f"🎯 Processing element: {element.name}, classes: {element_classes}")
                    
                    # Section titles (prioritize specific classes)
                    if 'section-title' in element_classes:
                        logger.info("✅ Processing section-title")
                        self._add_section_title(doc, element)
                    
                    # Report header container (only add once)
                    elif 'report-header-container' in element_classes:
                        if not self._header_added:
                            logger.info("✅ Processing report-header-container")
                            self._add_report_header(doc, element)
                            self._header_added = True
                            logger.info("✅ Added report header (first slide only)")
                        else:
                            logger.info("🔄 Skipping duplicate report header")
                    
                    # Company header with logo and rating
                    elif 'company-header' in element_classes:
                        logger.info("✅ Processing company-header")
                        # Check if this company header is inside a header-blue-border
                        parent = element.parent
                        parent_classes = parent.get('class', []) if parent else []
                        if 'header-blue-border' in parent_classes:
                            logger.info("🔵 Company header is inside header-blue-border, adding with border")
                            self._add_company_header_with_border(doc, element)
                        else:
                            # Check grandparent as well for nested structures
                            grandparent = parent.parent if parent else None
                            grandparent_classes = grandparent.get('class', []) if grandparent else []
                            if 'header-blue-border' in grandparent_classes:
                                logger.info("🔵 Company header grandparent has header-blue-border, adding with border")
                                self._add_company_header_with_border(doc, element)
                            else:
                                logger.info("🏢 Company header standalone, adding without border")
                                self._add_company_header(doc, element)
                    
                    # Metrics grid
                    elif 'metrics-grid' in element_classes:
                        logger.info("✅ Processing metrics-grid")
                        self._add_metrics_grid(doc, element)
                    
                    # Introduction and chart container
                    elif 'intro-and-chart-container' in element_classes:
                        logger.info("✅ Processing intro-and-chart-container")
                        self._add_intro_chart_container(doc, element)
                    
                    # Main content container (contains metrics, text, etc.)
                    elif element.name == 'main':
                        logger.info("✅ Processing main content container")
                        self._process_main_content(doc, element, slide_classes)
                    
                    # Analysis items  
                    elif 'analysis-item' in element.get('class', []):
                        self._add_analysis_item(doc, element)
                    
                    # Headers (h1, h2, h3, etc.)
                    elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                        self._add_header(doc, element)
                    
                    # Tables
                    elif element.name == 'table':
                        self._add_table(doc, element)
                    
                    # Slide logo container (skip - already processed at slide start)
                    elif 'slide-logo' in element_classes:
                        logger.info("🔄 Skipping slide-logo container (already processed at slide start)")
                        pass
                    
                    # Header blue border (contains company header)
                    elif 'header-blue-border' in element_classes:
                        logger.info("✅ Processing header-blue-border")
                        company_header = element.find(class_='company-header')
                        logger.info(f"🔍 DEBUG: Company header found: {bool(company_header)}")
                        if company_header:
                            logger.info("🔵 Calling _add_company_header_with_border")
                            self._add_company_header_with_border(doc, company_header)
                        else:
                            # Fallback: Look for any header-like content within
                            header_content = element.find(['h1', 'h2', 'h3'])
                            if header_content:
                                logger.info("🔵 Fallback: Found header content in header-blue-border, adding with border")
                                self._add_company_header_with_border(doc, element)
                            else:
                                logger.warning("⚠️ No company-header found in header-blue-border")
                    
                    # Regular paragraphs and divs (catch-all for content)
                    elif element.name in ['p', 'div', 'span']:
                        # Only add if it has meaningful text content
                        text_content = element.get_text().strip()
                        if text_content and len(text_content) > 0:
                            self._add_paragraph(doc, element)
                    
                    # Lists
                    elif element.name in ['ul', 'ol']:
                        self._add_list(doc, element)
            
            logger.info(f"✅ Processed {elements_processed} elements in slide content")
        else:
            logger.warning("⚠️ content_soup has no children attribute, attempting text extraction")
            # Fallback: extract text directly
            text_content = str(content_soup) if content_soup else ""
            if text_content.strip():
                p = doc.add_paragraph(text_content.strip())
                logger.info(f"✅ Added fallback paragraph with {len(text_content)} characters")
    
    def _add_section_title(self, doc: Document, element):
        """Add section title with Robeco styling and proper alignment"""
        p = doc.add_paragraph()
        # Use intelligent alignment detection
        self._fix_element_alignment(p, element)
        run = p.add_run(element.get_text().strip())
        run.font.size = Pt(26)
        run.font.name = self.primary_font
        run.font.bold = True
        run.font.color.rgb = self.robeco_colors['blue_darker']
        p.space_after = Pt(18)
        
        # Add underline border (approximate CSS border-bottom)
        p_pr = p._element.get_or_add_pPr()
        p_bdr = OxmlElement('w:pBdr')
        bottom_border = OxmlElement('w:bottom')
        bottom_border.set(qn('w:val'), 'single')
        bottom_border.set(qn('w:sz'), '30')  # 5px equivalent
        bottom_border.set(qn('w:color'), '005F90')  # Robeco blue
        p_bdr.append(bottom_border)
        p_pr.append(p_bdr)
    
    def _add_metrics_grid(self, doc: Document, grid_soup):
        """Add metrics grid as 5×5 table matching HTML CSS Grid structure"""
        metrics_items = grid_soup.find_all(class_='metrics-item')
        if not metrics_items:
            return
        
        logger.info(f"📊 Creating metrics grid with {len(metrics_items)} items (target: 5×5 structure)")
        
        # Create EXACTLY 5-column table (CSS: grid-template-columns: repeat(5, 1fr))
        # Calculate rows needed, ensuring we handle exactly 25 items for 5×5 grid
        rows_needed = max(5, (len(metrics_items) + 4) // 5)  # Ensure at least 5 rows for 5×5
        table = doc.add_table(rows=rows_needed, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set EQUAL column widths (CSS: repeat(5, 1fr) = equal fractions)
        table.autofit = False
        col_width = Inches(1.5)  # Each column = 1.5 inches (7.5 / 5 = 1.5)
        
        for col in table.columns:
            col.width = col_width
        
        # Set consistent row heights
        row_height = Inches(0.8)
        for row in table.rows:
            row.height = row_height
        
        # Add top border (CSS: border-top: 5px solid var(--robeco-blue))
        self._add_table_border(table, 'top')
        
        # Fill table with metrics in 5×5 grid pattern
        for i, item in enumerate(metrics_items):
            if i >= 25:  # Prevent overflow beyond 5×5 grid
                logger.warning(f"⚠️ More than 25 metrics items found, truncating at 25")
                break
                
            row_idx = i // 5
            col_idx = i % 5
            cell = table.cell(row_idx, col_idx)
            
            # Get label and value (HTML structure: .label and .value divs)
            label_elem = item.find(class_='label')
            value_elem = item.find(class_='value') 
            
            if label_elem and value_elem:
                # Add label (CSS: font-size: 10pt; font-weight: 700)
                label_p = cell.paragraphs[0]
                label_p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                label_run = label_p.add_run(label_elem.get_text().strip())
                label_run.font.size = Pt(10)
                label_run.font.bold = True
                label_run.font.color.rgb = self.robeco_colors['text_secondary']
                
                # Add value (CSS: font-size: 14pt; font-weight: 400)
                value_p = cell.add_paragraph()
                value_p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                value_run = value_p.add_run(value_elem.get_text().strip())
                value_run.font.size = Pt(14)
                value_run.font.name = self.primary_font
                value_run.font.bold = True
                value_run.font.color.rgb = self.robeco_colors['text_dark']
                
                # Set cell vertical alignment
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        
        # Add bottom border (CSS: border-bottom: 5px solid var(--robeco-blue))
        self._add_table_border(table, 'bottom')
        
        logger.info(f"✅ Created {rows_needed}×5 metrics grid table with {min(len(metrics_items), 25)} items")
    
    def _add_analysis_item(self, doc: Document, item_soup):
        """Add analysis item with proper formatting - ENHANCED: Direct two-column processing"""
        logger.info("🎯 ENHANCED: Processing analysis item with direct two-column layout")
        self._add_two_column_analysis_item(doc, item_soup)
    
    def _add_paragraph(self, doc: Document, p_soup):
        """Add paragraph with text formatting and proper alignment"""
        text = p_soup.get_text().strip()
        if not text:
            return
        
        p = doc.add_paragraph()
        # Use intelligent alignment detection
        self._fix_element_alignment(p, p_soup)
        
        # Handle citations and formatting
        if hasattr(p_soup, 'children'):
            for child in p_soup.children:
                if isinstance(child, NavigableString):
                    run = p.add_run(str(child))
                    run.font.size = Pt(18)
                    run.font.name = self.primary_font
                elif hasattr(child, 'name'):
                    if child.name == 'strong' or child.name == 'b':
                        run = p.add_run(child.get_text())
                        run.font.bold = True
                        run.font.size = Pt(18)
                        run.font.name = self.primary_font
                    elif child.name == 'em' or child.name == 'i':
                        run = p.add_run(child.get_text())
                        run.font.italic = True
                        run.font.size = Pt(18)
                        run.font.name = self.primary_font
                    else:
                        run = p.add_run(child.get_text())
                        run.font.size = Pt(18)
                        run.font.name = self.primary_font
        else:
            run = p.add_run(text)
            run.font.size = Pt(18)
            run.font.name = self.primary_font
        
        p.space_after = Pt(12)
    
    def _add_table(self, doc: Document, table_soup):
        """Add HTML table to Word document"""
        rows = table_soup.find_all('tr')
        if not rows:
            return
        
        # Create Word table
        max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Fill table content
        for i, row_soup in enumerate(rows):
            cells = row_soup.find_all(['td', 'th'])
            for j, cell_soup in enumerate(cells):
                if j < max_cols:
                    cell = table.cell(i, j)
                    cell_p = cell.paragraphs[0]
                    run = cell_p.add_run(cell_soup.get_text().strip())
                    
                    # Header styling for th elements
                    if cell_soup.name == 'th':
                        run.font.bold = True
                        run.font.color.rgb = self.robeco_colors['blue_darker']
                    
                    run.font.size = Pt(14)
                    run.font.name = self.primary_font
    
    def _add_header(self, doc: Document, header_element):
        """Add header element (h1, h2, h3, etc.) with appropriate styling and alignment"""
        text = header_element.get_text().strip()
        if not text:
            return
            
        p = doc.add_paragraph()
        # Use intelligent alignment detection
        self._fix_element_alignment(p, header_element)
        run = p.add_run(text)
        run.font.name = self.primary_font
        run.font.bold = True
        run.font.color.rgb = self.robeco_colors['blue_darker']
        
        # Size based on header level
        header_level = int(header_element.name[1])  # h1 -> 1, h2 -> 2, etc.
        sizes = {1: 28, 2: 24, 3: 20, 4: 18, 5: 16, 6: 14}
        run.font.size = Pt(sizes.get(header_level, 18))
        
        p.space_after = Pt(16)
    
    def _add_list(self, doc: Document, list_element):
        """Add list (ul or ol) to Word document with enhanced formatting"""
        # Check if this is within a bullet-list-square container
        parent_classes = []
        current = list_element.parent
        while current and hasattr(current, 'get'):
            parent_classes.extend(current.get('class', []))
            current = getattr(current, 'parent', None)
        
        is_special_bullet = 'bullet-list-square' in parent_classes
        
        items = list_element.find_all('li')
        for item in items:
            text = item.get_text().strip()
            if text:
                p = doc.add_paragraph()
                # Use intelligent alignment for list items
                if is_special_bullet:
                    self._fix_element_alignment(p, item)
                else:
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                
                # Add bullet or number with special formatting for bullet-list-square
                if list_element.name == 'ul':
                    p.style = doc.styles['List Bullet']
                else:
                    p.style = doc.styles['List Number']
                
                # Process text with potential formatting (bold, italic)
                if is_special_bullet:
                    # Use formatted text processing for special lists
                    self._add_formatted_text_to_paragraph(p, item)
                else:
                    # Standard list formatting
                    run = p.add_run(text)
                    run.font.size = Pt(16)
                    run.font.name = self.primary_font
                    run.font.color.rgb = self.robeco_colors['text_dark']
                
                # Add extra spacing for special bullet lists
                if is_special_bullet:
                    p.space_after = Pt(8)
    
    def _process_report_footer(self, doc: Document, footer_soup):
        """Add report footer"""
        footer_text = footer_soup.get_text().strip()
        if footer_text:
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run(footer_text)
            run.font.size = Pt(17)
            run.font.name = self.primary_font
            run.font.color.rgb = self.robeco_colors['text_secondary']
            
            # Add top border
            p_pr = p._element.get_or_add_pPr()
            p_bdr = OxmlElement('w:pBdr')
            top_border = OxmlElement('w:top')
            top_border.set(qn('w:val'), 'single')
            top_border.set(qn('w:sz'), '30')  # 5px equivalent
            top_border.set(qn('w:color'), '005F90')  # Robeco blue
            p_bdr.append(top_border)
            p_pr.append(p_bdr)
    
    def _add_table_border(self, table, position: str):
        """Add border to table (top or bottom)"""
        tbl_pr = table._element.find(qn('w:tblPr'))
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            table._element.insert(0, tbl_pr)
        
        tbl_borders = tbl_pr.find(qn('w:tblBorders'))
        if tbl_borders is None:
            tbl_borders = OxmlElement('w:tblBorders')
            tbl_pr.append(tbl_borders)
        
        border = OxmlElement(f'w:{position}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '30')  # 5px equivalent
        border.set(qn('w:color'), '005F90')  # Robeco blue
        tbl_borders.append(border)

    def _add_report_header(self, doc: Document, header_soup):
        """Add report header with proper Robeco styling and logos"""
        # Process Robeco logo if present - but ONLY once per document
        if not hasattr(self, '_robeco_logo_added'):
            self._robeco_logo_added = False
            
        if not self._robeco_logo_added:
            robeco_logo_container = header_soup.find(class_='robeco-logo-container')
            if robeco_logo_container:
                # Try to find actual Robeco logo image
                robeco_img = robeco_logo_container.find('img')
                if robeco_img:
                    img_src = robeco_img.get('src', '')
                    p = doc.add_paragraph()
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Robeco logo is positioned top-right
                    success = self._download_and_insert_image_inline(p, img_src, "Robeco Logo", Inches(1.2))
                    if not success:
                        # Fallback to text
                        run = p.add_run("ROBECO")
                        run.font.size = Pt(24)
                        run.font.name = self.primary_font
                        run.font.bold = True
                        run.font.color.rgb = self.robeco_colors['blue_darker']
                    p.space_after = Pt(12)
                    self._robeco_logo_added = True
                    logger.info("✅ Added Robeco logo to document header (once only)")
                else:
                    # No image, use text fallback
                    p = doc.add_paragraph("ROBECO")
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    run = p.runs[0]
                    run.font.size = Pt(24)
                    run.font.name = self.primary_font
                    run.font.bold = True
                    run.font.color.rgb = self.robeco_colors['blue_darker']
                    p.space_after = Pt(12)
                    self._robeco_logo_added = True
                    logger.info("✅ Added Robeco text logo to document header (once only)")
        else:
            logger.info("🔄 Skipping duplicate Robeco logo")
        
        # Process company header within the report header
        company_header = header_soup.find(class_='company-header')
        if company_header:
            self._add_company_header(doc, company_header)
        else:
            # Search for company header in broader scope
            parent = header_soup.parent if header_soup.parent else header_soup
            broader_company_header = parent.find(class_='company-header') if hasattr(parent, 'find') else None
            if broader_company_header:
                self._add_company_header(doc, broader_company_header)
    
    def _add_company_header(self, doc: Document, header_soup):
        """Add company header with proper flex layout: icon+name left, rating right"""
        logger.info("🏢 Adding company header with proper flex layout")
        
        # Create table to replicate CSS flex layout: display: flex, align-items: center, gap: 12px
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Set column widths for proper layout
        usable_width = Inches(7.5)  # Full page width minus margins
        left_width = Inches(5.0)    # Space for icon + name
        right_width = Inches(2.5)   # Space for rating
        
        table.columns[0].width = left_width
        table.columns[1].width = right_width
        
        # Hide table borders
        self._hide_table_borders(table)
        
        # LEFT CELL: Company icon + name
        left_cell = table.cell(0, 0)
        left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Add company icon and name together
        self._add_company_icon_to_cell(left_cell, header_soup)
        
        # RIGHT CELL: Investment rating
        right_cell = table.cell(0, 1)
        right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        rating_elem = header_soup.find(class_='rating')
        if rating_elem:
            rating_p = right_cell.paragraphs[0]
            rating_p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Align to right
            rating_run = rating_p.add_run(rating_elem.get_text().strip())
            rating_run.font.size = Pt(20)
            rating_run.font.name = self.primary_font
            rating_run.font.bold = True
            
            # Color based on rating text
            rating_text = rating_elem.get_text().strip().upper()
            if 'UNDERWEIGHT' in rating_text or 'UNDERPERFORM' in rating_text:
                rating_run.font.color.rgb = RGBColor(198, 40, 40)  # Red for negative rating
            elif 'OVERWEIGHT' in rating_text or 'OUTPERFORM' in rating_text:
                rating_run.font.color.rgb = RGBColor(76, 175, 80)  # Green for positive rating  
            else:
                rating_run.font.color.rgb = RGBColor(255, 152, 0)  # Orange for neutral rating
        
        # Add spacing after header
        spacing_para = doc.add_paragraph()
        spacing_para.space_after = Pt(16)
    
    def _add_company_icon_to_cell(self, cell, header_soup):
        """Add company icon and name to a table cell with proper inline layout"""
        try:
            # Look for company icon (Clearbit logo)
            img_tags = header_soup.find_all('img')
            
            for img in img_tags:
                img_src = img.get('src', '')
                img_alt = img.get('alt', '')
                img_classes = img.get('class', [])
                
                logger.info(f"📸 Found img: src='{img_src}', alt='{img_alt}', classes='{img_classes}'")
                
                # Check if it's a company icon (clearbit or has 'icon' class)
                if ('clearbit.com' in img_src or 'icon' in img_classes or 
                    'icon' in img_alt.lower() or img_src.endswith('.co.jp')):
                    
                    # Create inline company icon with name in the cell
                    company_name_elem = header_soup.find('h1', class_='name') or header_soup.find(class_='name')
                    if company_name_elem:
                        p = cell.paragraphs[0]
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        
                        # Generate intelligent fallback URLs
                        company_name = company_name_elem.get_text().strip()
                        fallback_urls = self._generate_intelligent_logo_fallbacks(img_src, company_name)
                        logger.info(f"🎯 Generated {len(fallback_urls)} intelligent logo fallbacks for '{company_name}'")
                        
                        # Try each fallback URL
                        icon_added = False
                        for i, fallback_url in enumerate(fallback_urls):
                            logger.info(f"📸 Downloading image: {fallback_url}")
                            success = self._download_and_insert_image_inline(p, fallback_url, img_alt, Inches(0.4))
                            if success:
                                logger.info(f"✅ Success with fallback URL {i}: {fallback_url}")
                                icon_added = True
                                break
                        
                        # Add company name next to icon
                        if icon_added:
                            # Add space between icon and name
                            p.add_run(" ")
                        
                        name_run = p.add_run(company_name)
                        name_run.font.size = Pt(20)
                        name_run.font.name = self.primary_font
                        name_run.font.bold = True
                        name_run.font.color.rgb = self.robeco_colors['text_dark']
                        
                        if icon_added:
                            logger.info(f"✅ Added company icon with name: {img_alt}")
                        else:
                            logger.info(f"⚠️ Added company name without icon: {company_name}")
                        return True
                        
        except Exception as e:
            logger.error(f"❌ Company icon insertion failed: {e}")
            return False
    
    def _add_intro_chart_container(self, doc: Document, container_soup):
        """Add introduction and chart container with exact 50/50 flexbox layout"""
        logger.info("🔍 Processing intro-and-chart-container (HTML: display: flex, gap: 30px)")
        logger.info(f"🔍 DEBUG: Container HTML preview: {str(container_soup)[:300]}...")
        logger.info(f"🔍 DEBUG: Container type: {type(container_soup)}")
        logger.info(f"🔍 DEBUG: Container tag name: {getattr(container_soup, 'name', 'None')}")
        
        # Extract content from both columns (HTML structure analysis)
        intro_block = container_soup.find(class_='intro-text-block')
        chart_area = container_soup.find(class_='stock-chart-container')
        logger.info(f"🔍 DEBUG: Direct stock-chart-container search result: {bool(chart_area)}")
        
        # ENHANCED CHART DETECTION: Look for various chart patterns
        if not chart_area:
            # Method 1: Inline chart containers with height/background styling
            chart_divs = container_soup.find_all('div', style=lambda x: x and 'height:' in x and 'background:' in x)
            logger.info(f"🔍 DEBUG: Found {len(chart_divs)} divs with height/background styling")
            if chart_divs:
                chart_area = chart_divs[0]
                logger.info("📊 Found inline chart container with height/background styling")
                logger.info(f"🔍 DEBUG: Chart area HTML preview: {str(chart_area)[:200]}...")
                logger.info(f"🔍 DEBUG: Chart area type: {type(chart_area)}")
                
                # Check for SVG in this chart area immediately
                svg_in_chart = chart_area.find('svg')
                logger.info(f"🔍 DEBUG: SVG found in chart area: {bool(svg_in_chart)}")
                if svg_in_chart:
                    logger.info(f"🔍 DEBUG: SVG details - viewBox: {svg_in_chart.get('viewBox')}, style: {svg_in_chart.get('style')}")
                    logger.info(f"🔍 DEBUG: SVG has {len(svg_in_chart.find_all('text'))} text elements")
            
            # Method 2: Any div containing chart-related keywords
            if not chart_area:
                chart_keywords = ['chart', 'stock', 'price', 'graph', 'visualization']
                for div in container_soup.find_all('div'):
                    div_classes = div.get('class', [])
                    div_text = div.get_text().lower()
                    if (any(keyword in ' '.join(div_classes).lower() for keyword in chart_keywords) or
                        any(keyword in div_text for keyword in ['stock', 'price', 'chart', 'hkd', '$'])):
                        chart_area = div
                        logger.info(f"📊 Found chart area by keyword detection: classes={div_classes}")
                        break
            
            # Method 3: Look for canvas/svg elements
            if not chart_area:
                canvas_svg = container_soup.find(['canvas', 'svg'])
                if canvas_svg:
                    chart_area = canvas_svg.parent or canvas_svg
                    logger.info("📊 Found chart area containing canvas/svg")
        
        # Method 4: If still no chart area, check if container itself contains chart data
        if not chart_area:
            container_text = container_soup.get_text().lower()
            if any(keyword in container_text for keyword in ['stock', 'price', 'hkd', '$', 'current:', 'range:']):
                chart_area = container_soup  # Use the whole container as chart area
                logger.info("📊 Using entire container as chart area (contains price data)")
        
        logger.info(f"📝 Found intro-text-block: {bool(intro_block)}")
        logger.info(f"📊 Found chart area: {bool(chart_area)}")
        
        # If neither intro nor chart found, this might not be the expected container
        if not intro_block and not chart_area:
            logger.warning("⚠️ Neither intro nor chart found - treating entire container as chart content")
            chart_area = container_soup
        
        # Create 2-column table to replicate CSS flexbox (flex: 1 + flex: 1 = 50/50)
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set EXACT 50/50 column widths (CSS: both elements have flex: 1)
        # Account for 30px gap (CSS: gap: 30px) = ~0.4 inches
        col_width = Inches(3.55)  # Exact 50/50 split = 3.55 inches each (7.5 - 0.4) / 2
        
        table.columns[0].width = col_width  # intro-text-block
        table.columns[1].width = col_width  # chart container
        
        logger.info(f"📐 Set exact 50/50 column widths: {col_width} inches each")
        
        # Configure table borders (hidden to match CSS)
        self._hide_table_borders(table)
        
        # LEFT COLUMN: Process intro text block
        left_cell = table.cell(0, 0)
        left_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        
        if intro_block:
            logger.info("✅ Processing intro-text-block in left column")
            # Clear default paragraph
            left_cell.paragraphs[0].clear()
            
            # Check if there are <p> tags or direct content
            paragraphs = intro_block.find_all('p')
            if paragraphs:
                # Process each paragraph in intro block
                for para in paragraphs:
                    cell_para = left_cell.add_paragraph()
                    self._fix_element_alignment(cell_para, para)
            else:
                # Handle direct content in the intro block
                logger.info("📝 Processing direct content in intro-text-block")
                cell_para = left_cell.add_paragraph()
                cell_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                
                # Process the HTML content with formatting
                self._add_html_formatted_text_to_paragraph(cell_para, intro_block)
                cell_para.space_after = Pt(12)
        
        # RIGHT COLUMN: Process chart area (if exists)
        right_cell = table.cell(0, 1)
        right_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        
        if chart_area:
            logger.info("✅ Processing chart area in right column")
            logger.info(f"🔍 DEBUG: About to call _add_simple_static_chart with chart_area type: {type(chart_area)}")
            logger.info(f"🔍 DEBUG: Chart area tag: {getattr(chart_area, 'name', 'None')}")
            logger.info(f"🔍 DEBUG: Chart area classes: {chart_area.get('class', [])}")
            
            # Check SVG one more time before calling the chart method
            final_svg_check = chart_area.find('svg')
            logger.info(f"🔍 DEBUG: Final SVG check before chart processing: {bool(final_svg_check)}")
            
            # Clear default paragraph
            right_cell.paragraphs[0].clear()
            
            # Add simple static stock chart content
            self._add_simple_static_chart(right_cell, chart_area)
        else:
            # No chart area found, make left column span full width
            logger.info("⚠️ No chart area found, using single column layout")
            # Clear right cell
            right_cell.paragraphs[0].clear()
        
        # Add spacing after container
        spacing_para = doc.add_paragraph()
        spacing_para.space_after = Pt(20)
    
    def _add_simple_static_chart(self, cell, chart_area):
        """Add stock chart content - converts SVG directly to image for Word"""
        logger.info("📊 === STARTING CHART PROCESSING ===")
        try:
            logger.info("📊 Processing chart area for stock chart content")
            logger.info(f"📊 DEBUG: Chart area HTML: {str(chart_area)[:500]}...")
            logger.info(f"📊 DEBUG: Chart area full classes: {chart_area.get('class', [])}")
            logger.info(f"📊 DEBUG: Chart area id: {chart_area.get('id', 'None')}")
            logger.info(f"📊 DEBUG: Chart area style: {chart_area.get('style', 'None')}")
            
            # Count all elements in chart area
            all_elements = chart_area.find_all()
            logger.info(f"📊 DEBUG: Total elements in chart area: {len(all_elements)}")
            
            # Look for specific elements
            h4_elements = chart_area.find_all('h4')
            svg_elements = chart_area.find_all('svg')
            div_elements = chart_area.find_all('div')
            logger.info(f"📊 DEBUG: Elements found - h4: {len(h4_elements)}, svg: {len(svg_elements)}, div: {len(div_elements)}")
            
            # First try to convert SVG directly to image
            svg_chart = chart_area.find('svg')
            logger.info(f"🔍 DEBUG: SVG search result: {bool(svg_chart)}")
            logger.info(f"🔍 DEBUG: Chart area type: {type(chart_area)}")
            logger.info(f"🔍 DEBUG: Chart area tag: {getattr(chart_area, 'name', 'None')}")
            
            if svg_chart:
                logger.info("📊 Found SVG chart - converting to image")
                image_success = self._convert_svg_to_image(cell, chart_area, svg_chart)
                if image_success:
                    logger.info("✅ Successfully converted SVG chart to image")
                    return
                else:
                    logger.info("⚠️ SVG conversion failed, falling back to text extraction")
            else:
                logger.info("⚠️ No SVG found, using fallback text extraction")
            
            # Fallback: Extract text data if SVG conversion fails
            current_price = None
            chart_title = None
            price_range = None
            
            # Method 1: Look for title elements
            title_elem = chart_area.find(class_='chart-title')
            logger.info(f"🔍 DEBUG: Title element found: {bool(title_elem)}")
            if title_elem:
                chart_title = title_elem.get_text().strip()
                logger.info(f"📈 Found chart title: {chart_title}")
            
            # Method 2: Look for current price elements  
            price_elem = chart_area.find(class_='current-price')
            logger.info(f"🔍 DEBUG: Price element found: {bool(price_elem)}")
            if price_elem:
                price_text = price_elem.get_text().strip()
                logger.info(f"💰 DEBUG: Price text: '{price_text}'")
                import re
                price_match = re.search(r'[\$HKD\s]*(\d+\.?\d*)', price_text)
                logger.info(f"💰 DEBUG: Price regex match: {bool(price_match)}")
                if price_match:
                    current_price = float(price_match.group(1))
                    logger.info(f"💰 Found current price: ${current_price:.2f}")
            
            # Method 3: Look for price range elements
            range_elem = chart_area.find(class_='price-range')
            logger.info(f"🔍 DEBUG: Range element found: {bool(range_elem)}")
            if range_elem:
                price_range = range_elem.get_text().strip()
                logger.info(f"📊 Found price range: {price_range}")
            
            # Method 4: Extract from general text content if no specific elements
            if not current_price:
                chart_text = chart_area.get_text()
                import re
                # Look for patterns like "Current: HKD $40.38", "$40.38", "HKD 40.38"
                price_matches = re.findall(r'(?:Current|Price|HKD|[\$])\s*[\$]?\s*(\d+\.?\d*)', chart_text)
                if price_matches:
                    current_price = float(price_matches[0])
                    logger.info(f"💰 Extracted price from text: ${current_price:.2f}")
                
                # Look for range patterns
                range_matches = re.findall(r'Range[:\s]*[\$]?\s*(\d+\.?\d*)\s*[-–]\s*[\$]?\s*(\d+\.?\d*)', chart_text)
                if range_matches:
                    min_price, max_price = range_matches[0]
                    price_range = f"${min_price} - ${max_price}"
                    logger.info(f"📊 Extracted range from text: {price_range}")
            
            # Look for SVG charts (original logic) if no simple data found
            if not current_price:
                svg_chart = chart_area.find('svg')
                if svg_chart:
                    logger.info("📊 Found SVG chart - extracting price data")
                    
                    # Extract chart title from h4 element
                    title_elem = chart_area.find('h4')
                    if title_elem:
                        chart_title = title_elem.get_text().strip()
                        logger.info(f"📈 Chart title: {chart_title}")
                        
                        # Extract current price from title (e.g., "5-Year Stock Price (BUOU.SI) - Current: S$0.95")
                        import re
                        title_price_match = re.search(r'Current:\s*S?\$?(\d+\.\d+)', chart_title)
                        if title_price_match:
                            current_price = float(title_price_match.group(1))
                            logger.info(f"💰 Extracted current price from title: S${current_price:.2f}")
                    
                    # Extract price data from SVG text elements (for range)
                    price_texts = svg_chart.find_all('text')
                    prices = []
                    
                    import re
                    for text_elem in price_texts:
                        text_content = text_elem.get_text().strip()
                        # Look for price patterns like "S$1.20", "$1.20", etc.
                        price_match = re.search(r'[S$]?\$?(\d+\.\d+)', text_content)
                        if price_match:
                            try:
                                price_val = float(price_match.group(1))
                                prices.append(price_val)
                                logger.info(f"📊 Found price in SVG: S${price_val:.2f}")
                            except ValueError:
                                pass
                    
                    if prices:
                        # If we didn't get current price from title, use max price
                        if not current_price:
                            current_price = max(prices)
                        min_price = min(prices)
                        max_price = max(prices)
                        price_range = f"S${min_price:.2f} - S${max_price:.2f}"
                        logger.info(f"📊 Extracted SVG price data: Current=S${current_price:.2f}, Range={price_range}")
            
            # Fallback: Look for JavaScript data
            if not current_price:
                script_tags = chart_area.find_all('script')
                for script in script_tags:
                    script_text = script.get_text()
                    if 'stockData' in script_text:
                        import re
                        price_matches = re.findall(r"'price':\s*([\d.]+)", script_text)
                        if price_matches:
                            current_price = float(price_matches[-1])
                            start_price = float(price_matches[0])
                            min_price = min(float(p) for p in price_matches)
                            max_price = max(float(p) for p in price_matches)
                            price_change = current_price - start_price
                            price_change_pct = (price_change / start_price) * 100
                            logger.info("📊 Extracted JavaScript chart data")
                            break
            
            logger.info(f"🔍 DEBUG: Final values - current_price={current_price}, chart_title='{chart_title}'")
            
            if current_price or chart_title:
                logger.info("✅ Chart data found - rendering chart content")
                # Chart title (use extracted title or default)
                title_para = cell.add_paragraph()
                title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                display_title = chart_title if chart_title else "📈 Stock Price Chart"
                title_run = title_para.add_run(display_title)
                title_run.font.size = Pt(14)
                title_run.font.bold = True
                title_run.font.name = self.primary_font
                title_run.font.color.rgb = self.robeco_colors['brown_black']
                title_para.space_after = Pt(10)
                
                if current_price:
                    # Current price (large) - detect currency from price range
                    currency = "S$" if "S$" in str(price_range) else "$"
                    price_para = cell.add_paragraph()
                    price_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    price_run = price_para.add_run(f"{currency}{current_price:.2f}")
                    price_run.font.size = Pt(20)
                    price_run.font.bold = True
                    price_run.font.name = self.primary_font
                    price_run.font.color.rgb = self.robeco_colors['blue']
                    price_para.space_after = Pt(8)
                    
                    # Price range summary
                    if price_range:
                        range_para = cell.add_paragraph()
                        range_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        range_run = range_para.add_run(f"Range: {price_range}")
                        range_run.font.size = Pt(12)
                        range_run.font.name = self.primary_font
                        range_run.font.color.rgb = self.robeco_colors['text_secondary']
                        range_para.space_after = Pt(8)
                
                    currency = "S$" if "S$" in str(price_range) else "$"
                    logger.info(f"✅ Added SVG chart content: Current {currency}{current_price:.2f}")
                else:
                    # Just show the chart title if no price data
                    logger.info("📊 Added chart title without price data")
            else:
                # No chart data found - add placeholder
                placeholder_para = cell.add_paragraph()
                placeholder_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                placeholder_run = placeholder_para.add_run("📊 [Stock Chart]")
                placeholder_run.font.size = Pt(14)
                placeholder_run.font.italic = True
                placeholder_run.font.color.rgb = self.robeco_colors['text_secondary']
                logger.info("⚠️ Added chart placeholder (no data found)")
                
        except Exception as e:
            logger.error(f"❌ Chart processing failed: {e}")
            # Add simple fallback
            fallback_para = cell.add_paragraph()
            fallback_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            fallback_run = fallback_para.add_run("📊 [Chart Unavailable]")
            fallback_run.font.size = Pt(12)
            fallback_run.font.italic = True
    
    def _add_standalone_chart(self, doc: Document, chart_element):
        """Add standalone chart element (SVG, canvas, etc.) directly to document"""
        try:
            logger.info(f"📊 Processing standalone chart: {chart_element.name}")
            
            # Add chart as a centered element
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Try to extract chart data if it's an SVG
            if chart_element.name == 'svg':
                # Look for text elements in SVG that might contain prices
                price_texts = chart_element.find_all('text')
                prices = []
                chart_title = None
                
                for text_elem in price_texts:
                    text_content = text_elem.get_text().strip()
                    # Look for price patterns
                    import re
                    price_match = re.search(r'[\$￥¥]?(\d+\.?\d*)', text_content)
                    if price_match:
                        try:
                            prices.append(float(price_match.group(1)))
                        except ValueError:
                            pass
                    
                    # Look for chart title patterns
                    if any(keyword in text_content.lower() for keyword in ['stock', 'price', 'chart', 'year', 'month']):
                        if not chart_title or len(text_content) > len(chart_title):
                            chart_title = text_content
                
                # Also check parent/sibling elements for context
                parent = chart_element.parent
                if parent:
                    # Look for nearby text that might contain chart info
                    nearby_texts = parent.find_all(text=True)
                    for text in nearby_texts:
                        text_content = text.strip()
                        price_match = re.search(r'[\$￥¥]?(\d+\.?\d*)', text_content)
                        if price_match:
                            try:
                                prices.append(float(price_match.group(1)))
                            except ValueError:
                                pass
                
                if prices:
                    current_price = max(prices)
                    title_text = chart_title if chart_title else "5-Year Stock Price"
                    
                    # Create a visual chart representation
                    self._add_visual_stock_chart(p, title_text, current_price, prices)
                    logger.info(f"✅ Added visual stock chart: {title_text} with price ${current_price:.2f}")
                else:
                    title_text = chart_title if chart_title else "Interactive Stock Chart"
                    
                    # Create a simple chart placeholder with visual elements
                    self._add_chart_placeholder(p, title_text)
                    logger.info(f"✅ Added visual chart placeholder: {title_text}")
            else:
                # Canvas or other chart element
                run = p.add_run("📊 Interactive Chart Element")
                run.font.size = Pt(14)
                run.font.color.rgb = self.robeco_colors['blue']
                logger.info(f"✅ Added standalone {chart_element.name} chart")
                
        except Exception as e:
            logger.error(f"❌ Standalone chart processing failed: {e}")
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run("📊 [Chart Element]")
            run.font.size = Pt(12)
            run.font.italic = True
            run.font.color.rgb = self.robeco_colors['text_secondary']
    
    def _add_company_header_with_border(self, doc: Document, header_soup):
        """Add company header with blue border styling"""
        logger.info("🔵 Adding company header with blue border")
        
        # Add the company header first
        self._add_company_header(doc, header_soup)
        
        # Add blue border after the header with minimal spacing
        border_para = doc.add_paragraph()
        border_para.space_before = Pt(3)  # Minimal space before
        border_para.space_after = Pt(6)   # Reduced space after
        
        # Set paragraph formatting to reduce height
        pPr = border_para._element.get_or_add_pPr()
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '120')  # 6pt after
        spacing.set(qn('w:line'), '240')   # Single line spacing
        spacing.set(qn('w:lineRule'), 'auto')
        pPr.append(spacing)
        
        # Add blue bottom border using paragraph border
        p_bdr = OxmlElement('w:pBdr')
        bottom_border = OxmlElement('w:bottom')
        bottom_border.set(qn('w:val'), 'single')
        bottom_border.set(qn('w:sz'), '18')  # 3px equivalent 
        bottom_border.set(qn('w:color'), '005F90')  # Robeco blue
        bottom_border.set(qn('w:space'), '0')  # No space from text
        p_bdr.append(bottom_border)
        pPr.append(p_bdr)
        
        logger.info("🔵 Added blue border below company header")
    
    def _add_visual_stock_chart(self, paragraph, title, current_price, prices):
        """Create a visual stock chart representation in Word"""
        try:
            # Chart title
            title_run = paragraph.add_run(f"📈 {title}\n")
            title_run.font.size = Pt(14)
            title_run.font.bold = True
            title_run.font.color.rgb = self.robeco_colors['blue']
            
            # Current price display
            price_run = paragraph.add_run(f"Current: ${current_price:.2f}\n")
            price_run.font.size = Pt(16)
            price_run.font.bold = True
            price_run.font.color.rgb = self.robeco_colors['accent_green'] if len(prices) > 1 and current_price >= min(prices) else self.robeco_colors['text_dark']
            
            # Create a simple ASCII chart representation
            if len(prices) > 1:
                min_price = min(prices)
                max_price = max(prices)
                price_range = max_price - min_price if max_price > min_price else 1
                
                # Create 5 data points for visualization
                chart_points = prices[:5] if len(prices) >= 5 else prices
                
                chart_line = ""
                for i, price in enumerate(chart_points):
                    if i == 0:
                        chart_line += "●"  # Start point
                    elif price > chart_points[i-1]:
                        chart_line += "↗"
                    elif price < chart_points[i-1]:
                        chart_line += "↘"
                    else:
                        chart_line += "→"
                    
                    if i < len(chart_points) - 1:
                        chart_line += "─"
                
                chart_run = paragraph.add_run(f"{chart_line}\n")
                chart_run.font.size = Pt(12)
                chart_run.font.color.rgb = self.robeco_colors['blue']
                
                # Price range info
                range_run = paragraph.add_run(f"Range: ${min_price:.2f} - ${max_price:.2f}")
                range_run.font.size = Pt(10)
                range_run.font.color.rgb = self.robeco_colors['text_secondary']
            
        except Exception as e:
            logger.error(f"❌ Visual chart creation failed: {e}")
            # Fallback to simple text
            fallback_run = paragraph.add_run(f"📈 {title} - ${current_price:.2f}")
            fallback_run.font.size = Pt(14)
            fallback_run.font.color.rgb = self.robeco_colors['blue']
    
    def _convert_svg_to_image(self, cell, chart_area, svg_chart):
        """Convert SVG chart directly to PNG image and insert into Word cell"""
        logger.info("🖼️ === ENTERING SVG TO IMAGE CONVERSION ===")
        logger.info(f"🖼️ DEBUG: SVG chart type: {type(svg_chart)}")
        logger.info(f"🖼️ DEBUG: SVG chart tag: {getattr(svg_chart, 'name', 'None')}")
        logger.info(f"🖼️ DEBUG: SVG viewBox: {svg_chart.get('viewBox')}")
        logger.info(f"🖼️ DEBUG: SVG style: {svg_chart.get('style')}")
        logger.info(f"🖼️ DEBUG: SVG HTML preview: {str(svg_chart)[:300]}...")
        
        # Use Puppeteer to capture the chart as image
        try:
            import tempfile
            import subprocess
            import os
            
            logger.info("📊 Converting chart to image using Puppeteer")
            
            # Extract the chart area HTML
            chart_html = str(chart_area)
            
            # Create a temporary HTML file with just the chart
            html_content = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        body {{ 
            margin: 0; 
            padding: 20px; 
            font-family: Arial, sans-serif;
            background: white;
        }}
    </style>
</head>
<body>
    {chart_html}
</body>
</html>
"""
            
            with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as temp_html:
                temp_html.write(html_content)
                temp_html_path = temp_html.name
            
            # Create temporary PNG file
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_png:
                temp_png_path = temp_png.name
            
            # Use Puppeteer to capture screenshot
            puppeteer_script = f'''
const puppeteer = require('puppeteer');
(async () => {{
    const browser = await puppeteer.launch({{headless: true}});
    const page = await browser.newPage();
    await page.goto('file://{temp_html_path}');
    await page.setViewport({{width: 500, height: 450}});
    await page.screenshot({{path: '{temp_png_path}', fullPage: true}});
    await browser.close();
}})();
'''
            
            with tempfile.NamedTemporaryFile(mode='w', suffix='.js', delete=False) as temp_js:
                temp_js.write(puppeteer_script)
                temp_js_path = temp_js.name
            
            # Run Puppeteer
            result = subprocess.run(['node', temp_js_path], 
                                  capture_output=True, text=True, timeout=30)
            
            if result.returncode == 0 and os.path.exists(temp_png_path):
                # Insert the image into Word
                img_para = cell.add_paragraph()
                img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                run = img_para.add_run()
                run.add_picture(temp_png_path, width=Inches(4.0))
                
                logger.info("✅ Successfully converted chart to image using Puppeteer")
                success = True
            else:
                logger.error(f"❌ Puppeteer failed: {result.stderr}")
                success = False
            
            # Clean up temporary files
            for temp_file in [temp_html_path, temp_png_path, temp_js_path]:
                if os.path.exists(temp_file):
                    os.unlink(temp_file)
            
            return success
            
        except ImportError as e:
            logger.warning(f"⚠️ SVG conversion libraries not available: {e}")
            logger.info("💡 Install with: pip install cairosvg pillow")
            return False
        except Exception as e:
            logger.error(f"❌ SVG to image conversion failed: {e}")
            return False
    
    def _add_chart_placeholder(self, paragraph, title):
        """Create a visual chart placeholder"""
        try:
            # Chart title
            title_run = paragraph.add_run(f"📊 {title}\n")
            title_run.font.size = Pt(14)
            title_run.font.bold = True
            title_run.font.color.rgb = self.robeco_colors['blue']
            
            # Visual placeholder
            placeholder_run = paragraph.add_run("●─●─●─●─● [Interactive Chart]\n")
            placeholder_run.font.size = Pt(12)
            placeholder_run.font.color.rgb = self.robeco_colors['text_secondary']
            
            # Note
            note_run = paragraph.add_run("View full chart in digital report")
            note_run.font.size = Pt(10)
            note_run.font.italic = True
            note_run.font.color.rgb = self.robeco_colors['text_secondary']
            
        except Exception as e:
            logger.error(f"❌ Chart placeholder creation failed: {e}")
            # Fallback
            fallback_run = paragraph.add_run(f"📊 {title}")
            fallback_run.font.size = Pt(14)
            fallback_run.font.color.rgb = self.robeco_colors['blue']
    
    def _add_intro_content(self, doc: Document, intro_block):
        """Add introduction/analysis content from intro-text-block"""
        try:
            logger.info("📝 Processing intro-text-block for analysis content")
            
            # Add a section header for analysis
            analysis_header = doc.add_paragraph()
            analysis_header.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            header_run = analysis_header.add_run("Investment Analysis")
            header_run.font.size = Pt(18)
            header_run.font.bold = True
            header_run.font.color.rgb = self.robeco_colors['blue']
            analysis_header.space_after = Pt(12)
            
            # Process all paragraphs and content in the intro block
            for child in intro_block.children:
                if hasattr(child, 'name') and child.name is not None:
                    if child.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                        logger.info(f"✅ Found header in intro: {child.name}")
                        self._add_header(doc, child)
                    elif child.name == 'p':
                        logger.info("✅ Found paragraph in intro")
                        self._add_paragraph(doc, child)
                    elif child.name == 'ul' or child.name == 'ol':
                        logger.info("✅ Found list in intro")
                        self._add_list(doc, child)
                    elif child.name == 'div':
                        # Process nested divs
                        div_text = child.get_text().strip()
                        if div_text:
                            logger.info(f"✅ Found content div in intro: {len(div_text)} chars")
                            self._add_paragraph(doc, child)
                elif hasattr(child, 'strip'):
                    # Text node
                    text_content = child.strip()
                    if text_content:
                        logger.info("✅ Found direct text in intro")
                        p = doc.add_paragraph(text_content)
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        for run in p.runs:
                            run.font.size = Pt(16)
                            run.font.name = self.primary_font
                            run.font.color.rgb = self.robeco_colors['text_dark']
            
            logger.info("✅ Successfully processed intro-text-block content")
            
        except Exception as e:
            logger.error(f"❌ Intro content processing failed: {e}")
            # Fallback - add whatever text we can extract
            text_content = intro_block.get_text().strip()
            if text_content:
                p = doc.add_paragraph(text_content)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    def _add_stock_chart_content(self, cell, chart_area):
        """Add sophisticated stock chart representation based on D3.js chart data"""
        try:
            # Look for stock data in script tags
            script_tags = chart_area.find_all('script')
            stock_data = None
            
            for script in script_tags:
                script_text = script.get_text()
                if 'stockData' in script_text and 'date' in script_text and 'price' in script_text:
                    # Extract stock data from JavaScript with enhanced parsing
                    import re
                    import json
                    
                    # Find the complete stockData array
                    match = re.search(r'stockData\s*=\s*(\[.*?\]);', script_text, re.DOTALL)
                    if match:
                        try:
                            # Convert JavaScript object notation to Python-readable JSON
                            data_str = match.group(1)
                            # Replace single quotes with double quotes for JSON compatibility
                            json_str = re.sub(r"'([^']*)':", r'"\1":', data_str)
                            json_str = re.sub(r":\s*'([^']*)'", r': "\1"', json_str)
                            
                            # Parse the complete dataset
                            data_list = json.loads(json_str)
                            stock_data = [(item['date'], float(item['price'])) for item in data_list]
                            break
                        except Exception as e:
                            logger.error(f"JSON parsing failed, using regex fallback: {e}")
                            # Fallback to regex extraction
                            price_matches = re.findall(r"'price':\s*([\d.]+)", data_str)
                            date_matches = re.findall(r"'date':\s*'([\d-]+)'", data_str)
                            if price_matches and date_matches:
                                stock_data = list(zip(date_matches, [float(p) for p in price_matches]))
                                break
            
            if stock_data and len(stock_data) > 0:
                # Add sophisticated chart representation
                title_para = cell.add_paragraph()
                title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                title_run = title_para.add_run("📈 Stock Price Performance (HKD)")
                title_run.font.size = Pt(14)
                title_run.font.bold = True
                title_run.font.name = self.primary_font
                title_run.font.color.rgb = self.robeco_colors['blue']
                title_para.space_after = Pt(8)
                
                # Calculate key metrics
                prices = [price for _, price in stock_data]
                start_price = prices[0]
                current_price = prices[-1]
                max_price = max(prices)
                min_price = min(prices)
                price_change = current_price - start_price
                price_change_pct = (price_change / start_price) * 100
                
                # Current price highlight
                current_para = cell.add_paragraph()
                current_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                current_run = current_para.add_run(f"Current Price: HKD ${current_price:.2f}")
                current_run.font.size = Pt(18)
                current_run.font.bold = True
                current_run.font.name = self.primary_font
                current_run.font.color.rgb = self.robeco_colors['blue']
                current_para.space_after = Pt(8)
                
                # Performance summary
                perf_para = cell.add_paragraph()
                perf_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                change_color = self.robeco_colors.get('accent_green', RGBColor(46, 125, 50)) if price_change >= 0 else self.robeco_colors.get('accent_red', RGBColor(198, 40, 40))
                change_symbol = "+" if price_change >= 0 else ""
                perf_run = perf_para.add_run(f"Period Change: {change_symbol}HKD ${price_change:.2f} ({change_symbol}{price_change_pct:.1f}%)")
                perf_run.font.size = Pt(12)
                perf_run.font.bold = True
                perf_run.font.name = self.primary_font
                perf_run.font.color.rgb = change_color
                perf_para.space_after = Pt(8)
                
                # Price range
                range_para = cell.add_paragraph()
                range_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                range_run = range_para.add_run(f"Range: HKD ${min_price:.2f} - ${max_price:.2f}")
                range_run.font.size = Pt(11)
                range_run.font.name = self.primary_font
                range_run.font.color.rgb = self.robeco_colors['text_secondary']
                range_para.space_after = Pt(12)
                
                # Simple ASCII-style chart representation
                chart_para = cell.add_paragraph()
                chart_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                # Create a simple trend visualization using characters
                trend_length = min(20, len(stock_data))
                if trend_length > 1:
                    # Sample data points for visualization
                    sample_indices = [int(i * (len(stock_data) - 1) / (trend_length - 1)) for i in range(trend_length)]
                    sample_prices = [prices[i] for i in sample_indices]
                    
                    # Normalize prices to create visual representation
                    price_range = max_price - min_price
                    if price_range > 0:
                        normalized = [(p - min_price) / price_range for p in sample_prices]
                        
                        # Create visual trend line using characters
                        trend_chars = []
                        for i, norm_price in enumerate(normalized):
                            if i == 0:
                                trend_chars.append("●")  # Start point
                            elif i == len(normalized) - 1:
                                trend_chars.append("●")  # End point
                            else:
                                # Use trend direction
                                if i > 0:
                                    if normalized[i] > normalized[i-1]:
                                        trend_chars.append("↗")
                                    elif normalized[i] < normalized[i-1]:
                                        trend_chars.append("↘")
                                    else:
                                        trend_chars.append("→")
                        
                        chart_run = chart_para.add_run(" ".join(trend_chars))
                        chart_run.font.size = Pt(12)
                        chart_run.font.name = self.primary_font
                        chart_run.font.color.rgb = self.robeco_colors['blue']
                
                # Recent key data points (last 6 months)
                data_para = cell.add_paragraph()
                data_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                data_run = data_para.add_run("Recent Data Points:")
                data_run.font.size = Pt(11)
                data_run.font.bold = True
                data_run.font.name = self.primary_font
                data_para.space_after = Pt(4)
                
                # Show last 6 data points
                recent_data = stock_data[-6:] if len(stock_data) >= 6 else stock_data
                for date, price in recent_data:
                    point_para = cell.add_paragraph()
                    point_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    point_text = f"• {date}: HKD ${price:.2f}"
                    point_run = point_para.add_run(point_text)
                    point_run.font.size = Pt(9)
                    point_run.font.name = self.primary_font
                    point_run.font.color.rgb = self.robeco_colors['text_secondary']
                    point_para.space_after = Pt(2)
                
                logger.info(f"✅ Added sophisticated stock chart with {len(stock_data)} data points, range: ${min_price:.2f}-${max_price:.2f}")
            else:
                # Enhanced fallback when no data found
                self._add_chart_fallback(cell)
                logger.info("⚠️ Added enhanced chart fallback (no data parsed)")
                
        except Exception as e:
            logger.error(f"❌ Stock chart processing failed: {e}")
            self._add_chart_fallback(cell)
    
    def _add_chart_fallback(self, cell):
        """Add enhanced chart fallback when data extraction fails"""
        title_para = cell.add_paragraph()
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title_para.add_run("📊 Stock Price Chart")
        title_run.font.size = Pt(14)
        title_run.font.bold = True
        title_run.font.name = self.primary_font
        title_run.font.color.rgb = self.robeco_colors['blue']
        
        desc_para = cell.add_paragraph()
        desc_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        desc_run = desc_para.add_run("Interactive D3.js chart available in web version")
        desc_run.font.size = Pt(10)
        desc_run.font.italic = True
        desc_run.font.name = self.primary_font
        desc_run.font.color.rgb = self.robeco_colors['text_secondary']
    
    def _add_analysis_paragraph(self, doc, analysis_div):
        """Add formatted analysis paragraph with proper styling"""
        try:
            text_content = analysis_div.get_text().strip()
            if text_content:
                para = doc.add_paragraph()
                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                run = para.add_run(text_content)
                run.font.size = Pt(11)
                run.font.name = self.primary_font
                run.font.color.rgb = self.robeco_colors.get('text_dark', RGBColor(0, 0, 0))
                logger.info(f"✅ Added analysis paragraph with {len(text_content)} characters")
        except Exception as e:
            logger.error(f"❌ Failed to add analysis paragraph: {e}")
    
    def _hide_table_borders(self, table):
        """Hide all table borders to mimic CSS borderless layout"""
        try:
            # Hide table borders
            for row in table.rows:
                for cell in row.cells:
                    # Get cell properties
                    tc_pr = cell._element.find(qn('w:tcPr'))
                    if tc_pr is None:
                        tc_pr = OxmlElement('w:tcPr')
                        cell._element.insert(0, tc_pr)
                    
                    # Create border element with no borders
                    tc_borders = OxmlElement('w:tcBorders')
                    for border_name in ['top', 'left', 'bottom', 'right']:
                        border = OxmlElement(f'w:{border_name}')
                        border.set(qn('w:val'), 'nil')
                        tc_borders.append(border)
                    
                    tc_pr.append(tc_borders)
        except Exception as e:
            logger.warning(f"⚠️ Could not hide table borders: {e}")
    
    def _process_paragraph_with_formatting(self, word_para, html_para):
        """Process HTML paragraph with proper bold/strong formatting preservation"""
        try:
            # Process all content elements in order
            for element in html_para.contents:
                if hasattr(element, 'name'):
                    if element.name == 'strong':
                        # Bold text
                        run = word_para.add_run(element.get_text())
                        run.font.bold = True
                        run.font.size = Pt(16)
                        run.font.name = self.primary_font
                    else:
                        # Other HTML tags - extract text
                        run = word_para.add_run(element.get_text())
                        run.font.size = Pt(16)
                        run.font.name = self.primary_font
                else:
                    # Plain text node
                    text = str(element).strip()
                    if text:
                        run = word_para.add_run(text)
                        run.font.size = Pt(16)
                        run.font.name = self.primary_font
        except Exception as e:
            # Fallback to simple text extraction
            run = word_para.add_run(html_para.get_text())
            run.font.size = Pt(16)
            run.font.name = self.primary_font

    def _process_main_content(self, doc: Document, main_soup, slide_classes=None):
        """Process main content container - handles both analysis-item layout and prose layout"""
        main_classes = main_soup.get('class', [])
        slide_classes = slide_classes or []
        logger.info(f"🔍 Processing MAIN content container with {len(list(main_soup.children))} children")
        logger.info(f"🔍 Main classes: {main_classes}")
        
        # Get slide ID for debugging
        slide_parent = main_soup.parent
        slide_id = slide_parent.get('id', 'no-id') if slide_parent else 'no-id'
        logger.info(f"🔍 Processing main for slide ID: {slide_id}")
        
        # Check if this is a prose layout (Pages 3-15) - check both main and slide classes
        # Pages 1-2 should use analysis layout (portrait-page-1, portrait-page-1A)
        is_page_1_or_2 = slide_id in ['portrait-page-1', 'portrait-page-1A']
        is_prose = 'report-prose' in main_classes or 'report-prose' in slide_classes
        
        # Force analysis layout for pages 1-2 regardless of classes
        if is_page_1_or_2:
            logger.info(f"📊 Processing ANALYSIS LAYOUT (Pages 1-2) - Slide ID: {slide_id}")
            self._process_analysis_layout(doc, main_soup)
        elif is_prose:
            logger.info("📝 Processing PROSE LAYOUT (Pages 3-15)")
            self._process_prose_layout(doc, main_soup)
        else:
            logger.info("📊 Processing ANALYSIS LAYOUT (default)")
            self._process_analysis_layout(doc, main_soup)
    
    def _process_prose_layout(self, doc: Document, main_soup):
        """Process prose layout for Pages 3-15 - FULL WIDTH content"""
        for child in main_soup.children:
            if hasattr(child, 'name') and child.name is not None:
                if child.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    logger.info(f"✅ Found prose header {child.name}")
                    self._add_header(doc, child)
                elif child.name == 'p':
                    logger.info("✅ Found prose paragraph")
                    self._add_paragraph(doc, child)
                elif child.name == 'table':
                    logger.info("✅ Found FULL-WIDTH table in prose layout")
                    self._add_full_width_table(doc, child)
                elif child.name == 'div':
                    child_classes = child.get('class', [])
                    if 'content-item' in child_classes:
                        logger.info("✅ Found content-item in prose layout")
                        # Process content-item as simple paragraphs
                        for para in child.find_all('p'):
                            self._add_paragraph(doc, para)
                    elif 'metrics-grid' in child_classes:
                        logger.info("✅ Found metrics-grid in prose layout")
                        self._add_metrics_grid(doc, child)
                    elif 'intro-and-chart-container' in child_classes:
                        logger.info("⚠️ Found intro-chart container in prose - treating as full-width")
                        # Don't use 2-column layout in prose, extract content directly
                        self._add_full_width_intro_content(doc, child)
                    else:
                        # FIRST: Check for charts in this div before recursive processing
                        svg_elements = child.find_all('svg')
                        canvas_elements = child.find_all('canvas')
                        
                        # Enhanced chart detection
                        is_chart_div = any(keyword in ' '.join(child_classes).lower() for keyword in ['chart', 'stock', 'price', 'graph'])
                        
                        # Check for chart-indicating height in style attribute
                        style_attr = child.get('style', '')
                        is_chart_by_height = ('height' in style_attr and 
                                             any(size in style_attr for size in ['300px', '400px', '420px', '500px'])) or \
                                            ('height:' in style_attr and 
                                             any(size in style_attr for size in ['300', '400', '420', '500']))
                        
                        # Check for SVG with viewBox (strong chart indicator)
                        has_viewbox = any(svg.get('viewbox') or svg.get('viewBox') for svg in svg_elements)
                        
                        # Additional chart indicators
                        has_chart_title = child.find('h4') and any(word in child.find('h4').get_text().lower() for word in ['price', 'chart', 'stock', 'performance'])
                        contains_background_white = 'background: white' in style_attr or 'background-color: white' in style_attr
                        
                        # Check for structured content (tables, lists, headers) FIRST
                        has_tables = child.find_all('table')
                        has_lists = child.find_all(['ul', 'ol'])
                        has_headers = child.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])
                        has_structured_content = bool(has_tables or has_lists or has_headers)
                        
                        text_content = child.get_text().strip()
                        is_analysis_text = (len(text_content) > 200 and 
                                           any(keyword in text_content.lower() for keyword in ['overweight', 'underweight', 'buy', 'sell', 'hold', 'target', 'consensus', 'analysis', 'opportunity', 'presents', 'compelling'])
                                           and not has_structured_content)  # Only treat as analysis if no structured content
                        
                        logger.info(f"🔍 RECURSIVE DIV: classes={child_classes}, SVG={len(svg_elements)}, Canvas={len(canvas_elements)}, Tables={len(has_tables)}, Is chart={is_chart_div}, Has viewBox={has_viewbox}, Chart height={is_chart_by_height}, Chart title={has_chart_title}, White bg={contains_background_white}, Is analysis={is_analysis_text}, Text len={len(text_content)}")
                        
                        # Process charts (enhanced detection with multiple indicators)
                        is_likely_chart = (svg_elements or canvas_elements or is_chart_div or 
                                         is_chart_by_height or has_viewbox or has_chart_title or
                                         (contains_background_white and len(svg_elements) > 0))
                        
                        if is_likely_chart:
                            logger.info(f"📊 RECURSIVE: Found chart content - processing visual elements")
                            
                            # Process SVG charts
                            for svg in svg_elements:
                                self._add_standalone_chart(doc, svg)
                            
                            # Process canvas charts
                            for canvas in canvas_elements:
                                self._add_standalone_chart(doc, canvas)
                                
                            # If no visual elements but looks like chart container, add placeholder
                            if not svg_elements and not canvas_elements and (is_chart_div or is_chart_by_height or has_chart_title):
                                logger.info("📊 RECURSIVE: Chart container but no SVG/Canvas - adding placeholder")
                                self._add_chart_placeholder(doc, child)
                        
                        elif has_structured_content:
                            logger.info(f"📊 RECURSIVE: Found structured content (tables={len(has_tables)}, lists={len(has_lists)}, headers={len(has_headers)}) - processing elements")
                            # Process structured content properly
                            for table in has_tables:
                                logger.info("📊 RECURSIVE: Processing table structure")
                                self._add_full_width_table(doc, table)
                            for header in has_headers:
                                logger.info("📊 RECURSIVE: Processing header")
                                self._add_header(doc, header)
                            for list_elem in has_lists:
                                logger.info("📊 RECURSIVE: Processing list")
                                self._add_list(doc, list_elem)
                            # Process any remaining paragraphs that aren't inside tables/lists
                            for para in child.find_all('p'):
                                if not para.find_parent(['table', 'ul', 'ol']):  # Don't duplicate content already in tables/lists
                                    self._add_paragraph(doc, para)
                        
                        elif is_analysis_text:
                            logger.info(f"📝 RECURSIVE: Found pure analysis text - adding formatted paragraph")
                            self._add_analysis_paragraph(doc, child)
                            # Skip further processing to avoid duplication
                        else:
                            logger.info(f"✅ Processing div with classes: {child_classes}")
                            # Process child content recursively
                            self._process_prose_layout(doc, child)
                else:
                    logger.info(f"✅ Processing other element: {child.name}")
                    # Process other elements recursively  
                    self._process_prose_layout(doc, child)
    
    def _process_analysis_layout(self, doc: Document, main_soup):
        """Process analysis layout for Pages 1-2 with metrics grid, intro-chart, and analysis sections"""
        # DEBUG: Save main content for analysis
        logger.info(f"🔍 DEBUG: Main HTML content preview: {str(main_soup)[:500]}...")
        
        # Recursively process all children of main element
        for child in main_soup.children:
            if hasattr(child, 'name') and child.name is not None:
                child_classes = child.get('class', [])
                logger.info(f"🎯 Processing MAIN child: {child.name}, classes: {child_classes}")
                
                # Handle specific content types within main
                if 'metrics-grid' in child_classes:
                    logger.info("✅ Found metrics-grid in main content")
                    self._add_metrics_grid(doc, child)
                
                elif 'intro-and-chart-container' in child_classes:
                    logger.info("✅ Found intro-and-chart-container in main content")
                    self._add_intro_chart_container(doc, child)
                
                elif 'analysis-sections' in child_classes:
                    logger.info("✅ Found analysis-sections in main content")
                    self._process_analysis_sections(doc, child)
                
                # DEBUG: Look for any chart-related containers or SVG elements
                elif any(keyword in ' '.join(child_classes).lower() for keyword in ['chart', 'stock', 'price', 'graph']):
                    logger.info(f"🔍 DEBUG: Found potential chart container: {child.name}, classes: {child_classes}")
                    # Try to process as chart container
                    self._add_intro_chart_container(doc, child)
                
                # Look for standalone chart elements (SVG, canvas, chart divs)
                elif child.name in ['svg', 'canvas'] or any(keyword in ' '.join(child_classes).lower() for keyword in ['visualization', 'd3']):
                    logger.info(f"📊 DEBUG: Found standalone chart element: {child.name}, classes: {child_classes}")
                    # Create a simple chart display from SVG/canvas
                    self._add_standalone_chart(doc, child)
                    
                # Also check for divs that might contain charts
                elif child.name == 'div':
                    # Check if this div contains chart-related content
                    svg_elements = child.find_all('svg')
                    canvas_elements = child.find_all('canvas')
                    chart_class_elements = child.find_all(class_=lambda x: x and any(kw in ' '.join(x).lower() for kw in ['chart', 'stock', 'price', 'graph']))
                    
                    # Also check for script tags with chart data
                    script_elements = child.find_all('script')
                    chart_scripts = [s for s in script_elements if 'stockData' in s.get_text() or 'chartData' in s.get_text()]
                    
                    logger.info(f"📊 DEBUG: Checking div for charts: SVG={len(svg_elements)}, Canvas={len(canvas_elements)}, Chart classes={len(chart_class_elements)}, Chart scripts={len(chart_scripts)}")
                    
                    if svg_elements or canvas_elements or chart_class_elements or chart_scripts:
                        logger.info(f"📊 DEBUG: Found div with chart content: {len(svg_elements + canvas_elements + chart_class_elements + chart_scripts)} chart elements")
                        # Process the chart elements directly
                        for svg in svg_elements:
                            self._add_standalone_chart(doc, svg)
                        for canvas in canvas_elements:
                            self._add_standalone_chart(doc, canvas)
                        if chart_scripts:
                            # If we have chart scripts, try to process as chart container
                            logger.info(f"📊 DEBUG: Processing chart scripts in div")
                            self._add_intro_chart_container(doc, child)
                        elif chart_class_elements and not svg_elements and not canvas_elements:
                            # If only chart classes but no SVG/canvas, try full container processing
                            self._add_intro_chart_container(doc, child)
                    else:
                        # Process as regular content
                        logger.info(f"✅ Found content div in main: {len(child.get_text())} chars, classes: {child.get('class', [])}")
                        # Add div content as paragraph
                        if child.get_text().strip():
                            self._add_paragraph(doc, child)
                
                elif 'content-grid' in child_classes:
                    logger.info("✅ Found content-grid in main content") 
                    self._process_content_grid(doc, child)
                
                elif 'analysis-item' in child_classes:
                    logger.info("✅ Found analysis-item in main content")
                    self._add_two_column_analysis_item(doc, child)
                
                elif 'section' in child_classes or child.name == 'section':
                    logger.info("✅ Found section in main content")
                    # Process section content recursively
                    self._process_section_content(doc, child)
                
                elif child.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    logger.info(f"✅ Found header {child.name} in main content")
                    self._add_header(doc, child)
                
                elif child.name == 'p':
                    logger.info("✅ Found paragraph in main content")
                    self._add_paragraph(doc, child)
                
                elif child.name == 'table':
                    logger.info("✅ Found table in main content")
                    self._add_table(doc, child)
                
                elif child.name in ['ul', 'ol']:
                    logger.info(f"✅ Found list {child.name} in main content")
                    self._add_list(doc, child)
                
                elif child.name == 'div':
                    # Enhanced div processing - check for specific layout classes
                    if 'analysis-sections' in child_classes:
                        logger.info("✅ Found analysis-sections in div")
                        self._process_analysis_sections(doc, child)
                    elif 'content-grid' in child_classes:
                        logger.info("✅ Found content-grid in div")
                        self._process_content_grid(doc, child)
                    elif 'analysis-item' in child_classes:
                        logger.info("✅ Found analysis-item in div")
                        self._add_two_column_analysis_item(doc, child)
                    elif 'bullet-list-square' in child_classes:
                        logger.info("✅ Found bullet-list-square in div")
                        self._process_bullet_list_square(doc, child)
                    else:
                        # Check if this could be analysis content or chart content
                        intro_text_block = child.find(class_='intro-text-block') 
                        stock_chart_container = child.find(class_='stock-chart-container')
                        has_svg_or_chart = bool(child.find('svg') or child.find('canvas') or 
                                               any('chart' in str(c).lower() for c in child.get('class', [])))
                        
                        if intro_text_block and stock_chart_container:
                            logger.info("🔍 Found implicit intro-and-chart-container")
                            self._add_intro_chart_container(doc, child)
                        elif intro_text_block:
                            logger.info("📝 Found intro-text-block - adding analysis content")
                            self._add_intro_content(doc, intro_text_block)
                        elif has_svg_or_chart:
                            logger.info("📊 Found chart content in div")
                            # Process chart elements
                            for svg in child.find_all('svg'):
                                self._add_standalone_chart(doc, svg)
                        else:
                            # Generic div - check if it has meaningful content
                            text_content = child.get_text().strip()
                            if text_content and len(text_content) > 0:
                                logger.info(f"✅ Found content div in main: {len(text_content)} chars, classes: {child_classes}")
                                # If it has nested structure, recurse first
                                if child.find_all(['div', 'p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'table']):
                                    logger.info("🔄 Recursing into structured div content")
                                    self._process_main_content(doc, child)
                                else:
                                    self._add_paragraph(doc, child)
                            else:
                                # Empty div, but may contain other elements - recurse
                                logger.info("🔄 Recursing into empty div in main content")
                                self._process_main_content(doc, child)
                
                else:
                    # Generic element with text content
                    text_content = child.get_text().strip()
                    if text_content and len(text_content) > 0:
                        logger.info(f"✅ Found generic content element {child.name}: {len(text_content)} chars")
                        self._add_paragraph(doc, child)

    def _process_section_content(self, doc: Document, section_soup):
        """Process section content with proper handling"""
        logger.info(f"🔍 Processing SECTION content")
        
        # Look for section title
        section_title = section_soup.find(class_='section-title')
        if section_title:
            self._add_section_title(doc, section_title)
        
        # Process other content in section
        for child in section_soup.children:
            if hasattr(child, 'name') and child.name is not None:
                if 'section-title' not in child.get('class', []):  # Skip already processed title
                    if child.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                        self._add_header(doc, child)
                    elif child.name == 'p':
                        self._add_paragraph(doc, child)
                    elif child.name == 'table':
                        self._add_table(doc, child)
                    elif child.name in ['ul', 'ol']:
                        self._add_list(doc, child)

    def _process_analysis_sections(self, doc: Document, sections_soup):
        """Process analysis sections with proper hierarchy and structure"""
        logger.info("🔍 Processing ANALYSIS SECTIONS with structured content")
        
        # Look for h3 headers first
        headers = sections_soup.find_all('h3')
        for header in headers:
            logger.info(f"📋 Found section header: {header.get_text().strip()}")
            self._add_header(doc, header)
        
        # Process paragraphs
        paragraphs = sections_soup.find_all('p')
        for para in paragraphs:
            logger.info(f"📝 Found section paragraph: {len(para.get_text())} chars")
            self._add_paragraph(doc, para)
        
        # Look for bullet lists
        bullet_lists = sections_soup.find_all(class_='bullet-list-square')
        for bullet_list in bullet_lists:
            logger.info("🔘 Found bullet-list-square in analysis section")
            self._process_bullet_list_square(doc, bullet_list)
        
        # Process any remaining analysis items
        analysis_items = sections_soup.find_all(class_='analysis-item')
        for item in analysis_items:
            logger.info("📊 Found analysis-item in analysis section")
            self._add_two_column_analysis_item(doc, item)

    def _process_content_grid(self, doc: Document, grid_soup):
        """Process content grid layouts"""
        logger.info("🔍 Processing CONTENT GRID layout")
        
        # Look for content blocks within the grid
        content_blocks = grid_soup.find_all(class_='content-block')
        
        for block in content_blocks:
            logger.info("📦 Processing content block within grid")
            
            # Process section titles
            section_titles = block.find_all(class_='section-title')
            for title in section_titles:
                self._add_section_title(doc, title)
            
            # Process other content recursively
            self._process_main_content(doc, block)

    def _add_two_column_analysis_item(self, doc: Document, item_soup):
        """Add two-column analysis item with PERFECT HTML flexbox replication (16%/84% layout)"""
        logger.info("🎯 ENHANCED: Processing TWO-COLUMN ANALYSIS ITEM with precise HTML structure replication")
        
        # 🔍 ANALYZE HTML STRUCTURE: Extract precise layout information
        item_title = item_soup.find(class_='item-title')
        content_item = item_soup.find(class_='content-item')
        
        # Get CSS classes for advanced layout detection
        item_classes = item_soup.get('class', [])
        is_first_item = 'first-analysis-item' in item_classes
        
        logger.info(f"🔍 Structure Analysis:")
        logger.info(f"   📋 Has item-title: {bool(item_title)}")
        logger.info(f"   📄 Has content-item: {bool(content_item)}")
        logger.info(f"   🎯 Is first item: {is_first_item}")
        logger.info(f"   📦 Item classes: {item_classes}")
        
        if item_title and content_item:
            title_text = item_title.get_text().strip()
            content_text = content_item.get_text().strip()
            
            logger.info(f"📋 Title: '{title_text}'")
            logger.info(f"📄 Content: {len(content_text)} chars")
            
            # 🏗️ CREATE PRECISE FLEXBOX-STYLE TABLE: Mimic CSS display:flex with flex:0 0 16% and flex:1
            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            
            # 🎨 APPLY EXACT HTML STYLING: Replicate CSS border-bottom and padding
            self._apply_analysis_item_html_styling(table, is_first_item)
            
            # 📐 FORCE EXACT 16%/84% COLUMN LAYOUT: Use advanced XML manipulation
            self._force_precise_flexbox_layout(table)
            
            # 📋 LEFT COLUMN (16%): Item Title - Exact HTML replication
            left_cell = table.cell(0, 0)
            left_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            
            # Apply exact HTML title styling (.item-title CSS)
            title_p = left_cell.paragraphs[0]
            title_p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            title_run = title_p.add_run(title_text)
            title_run.font.size = Pt(16)  # Matches CSS font-size: 16px
            title_run.font.name = self.primary_font
            title_run.font.bold = True  # Matches CSS font-weight: 500 (bold)
            title_run.font.color.rgb = self.robeco_colors['blue_darker']  # Matches CSS color: var(--robeco-blue)
            
            # Apply HTML padding-right: 15px equivalent
            self._apply_cell_padding(left_cell, right_padding=15)
            
            # 📄 RIGHT COLUMN (84%): Content Item - Perfect HTML content structure
            right_cell = table.cell(0, 1)
            right_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            
            # 🔍 DEEP CONTENT ANALYSIS: Process nested HTML structure precisely
            self._process_content_item_with_html_structure(right_cell, content_item)
            
            logger.info(f"✅ ENHANCED: Two-column analysis created with perfect HTML replication")
            
            # Add proper spacing (matches CSS margin-bottom)
            spacing_para = doc.add_paragraph()
            spacing_para.space_after = Pt(12)  # Matches HTML margin-bottom: 10px + padding
            
        else:
            logger.warning("⚠️ FALLBACK: Analysis item missing required structure - using basic processing")
            # Fallback to basic content processing
            if item_soup.get_text().strip():
                self._add_paragraph(doc, item_soup)
    
    def _apply_analysis_item_html_styling(self, table, is_first_item: bool):
        """Apply exact HTML analysis-item CSS styling to Word table"""
        logger.info(f"🎨 Applying HTML analysis-item styling (first_item: {is_first_item})")
        
        try:
            # Get table properties
            tbl_pr = table._element.find(qn('w:tblPr'))
            if tbl_pr is None:
                tbl_pr = OxmlElement('w:tblPr')
                table._element.insert(0, tbl_pr)
            
            # Add table borders to match HTML styling
            tbl_borders = tbl_pr.find(qn('w:tblBorders'))
            if tbl_borders is None:
                tbl_borders = OxmlElement('w:tblBorders')
                tbl_pr.append(tbl_borders)
            
            # Bottom border: matches CSS border-bottom: 3.5px solid #E0E0E0 EXACTLY
            bottom_border = OxmlElement('w:bottom')
            bottom_border.set(qn('w:val'), 'single')
            bottom_border.set(qn('w:sz'), '21')  # 3.5px equivalent
            bottom_border.set(qn('w:color'), 'E0E0E0')  # Light gray to match CSS exactly
            tbl_borders.append(bottom_border)
            
            # Top border for first item: matches CSS .first-analysis-item { border-top: 5px solid var(--robeco-blue) }
            if is_first_item:
                top_border = OxmlElement('w:top')
                top_border.set(qn('w:val'), 'single')
                top_border.set(qn('w:sz'), '30')  # 5px equivalent
                top_border.set(qn('w:color'), '005F90')  # Robeco blue
                tbl_borders.append(top_border)
                logger.info("🔵 Added top border for first analysis item")
            
            # Set table spacing to match HTML padding-top: 9px, padding-bottom: 9px
            tbl_cellmargin = OxmlElement('w:tblCellMar')
            
            # Top and bottom padding
            for direction, value in [('top', '54'), ('bottom', '54')]:  # 9px = 54 twips
                margin = OxmlElement(f'w:{direction}')
                margin.set(qn('w:w'), value)
                margin.set(qn('w:type'), 'dxa')
                tbl_cellmargin.append(margin)
            
            tbl_pr.append(tbl_cellmargin)
            
            logger.info("✅ Applied exact HTML analysis-item styling")
            
        except Exception as e:
            logger.error(f"❌ Error applying HTML styling: {e}")
    
    def _force_precise_flexbox_layout(self, table):
        """Force EXACT 16%/84% flexbox layout using percentage-based approach for better Word compatibility"""
        logger.info("📐 PRECISION: Forcing exact CSS flexbox layout (16%/84%) to match HTML")
        
        try:
            tbl = table._element
            
            # Set table properties for full width using percentage
            tblPr = tbl.find(qn('w:tblPr'))
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            
            # Table width: 100% (use percentage for better compatibility)
            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:w'), '5000')  # 100% width in percentage units
            tblW.set(qn('w:type'), 'pct')
            tblPr.append(tblW)
            
            # Force individual cell widths using percentages
            for row in tbl.findall(qn('w:tr')):
                cells = row.findall(qn('w:tc'))
                if len(cells) >= 2:
                    # Cell 1: EXACT 16% width
                    tcPr1 = cells[0].find(qn('w:tcPr'))
                    if tcPr1 is None:
                        tcPr1 = OxmlElement('w:tcPr')
                        cells[0].insert(0, tcPr1)
                    
                    tcW1 = OxmlElement('w:tcW')
                    tcW1.set(qn('w:w'), '800')  # 16% in percentage units (16% of 5000)
                    tcW1.set(qn('w:type'), 'pct')
                    tcPr1.append(tcW1)
                    
                    # Cell 2: EXACT 84% width  
                    tcPr2 = cells[1].find(qn('w:tcPr'))
                    if tcPr2 is None:
                        tcPr2 = OxmlElement('w:tcPr')
                        cells[1].insert(0, tcPr2)
                    
                    tcW2 = OxmlElement('w:tcW')
                    tcW2.set(qn('w:w'), '4200')  # 84% in percentage units (84% of 5000)
                    tcW2.set(qn('w:type'), 'pct')
                    tcPr2.append(tcW2)
            
            logger.info("✅ PRECISION: Applied EXACT 16%/84% layout using percentage-based approach")
            
        except Exception as e:
            logger.error(f"❌ PRECISION FAILED: {e}")
    
    def _apply_cell_padding(self, cell, right_padding: int = 0):
        """Apply cell padding to match HTML CSS padding"""
        try:
            if right_padding > 0:
                # Convert px to twips (1px = 6 twips)
                padding_twips = str(right_padding * 6)
                
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcMar = OxmlElement('w:tcMar')
                
                right_mar = OxmlElement('w:right')
                right_mar.set(qn('w:w'), padding_twips)
                right_mar.set(qn('w:type'), 'dxa')
                tcMar.append(right_mar)
                
                tcPr.append(tcMar)
                
                logger.info(f"✅ Applied right padding: {right_padding}px ({padding_twips} twips)")
        except Exception as e:
            logger.error(f"❌ Error applying cell padding: {e}")
    
    def _process_content_item_with_html_structure(self, cell, content_item):
        """Process content-item with exact HTML structure preservation"""
        logger.info("📄 PRECISION: Processing content-item with HTML structure analysis")
        
        # Analyze the content structure
        content_paragraphs = content_item.find_all('p')
        content_lists = content_item.find_all(['ul', 'ol'])
        content_text = content_item.get_text().strip()
        
        logger.info(f"📊 Content Analysis:")
        logger.info(f"   📝 Paragraphs: {len(content_paragraphs)}")
        logger.info(f"   📋 Lists: {len(content_lists)}")
        logger.info(f"   📄 Total text: {len(content_text)} chars")
        
        if content_paragraphs:
            # Process each paragraph with exact HTML formatting
            for i, para in enumerate(content_paragraphs):
                para_text = para.get_text().strip()
                
                if para_text:  # Only process non-empty paragraphs
                    if i == 0:
                        # Use existing cell paragraph
                        content_p = cell.paragraphs[0]
                    else:
                        # Add new paragraph
                        content_p = cell.add_paragraph()
                    
                    # Apply exact HTML content styling (.content-item CSS)
                    content_p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # Matches typical content alignment
                    
                    # Process with advanced text formatting that preserves HTML structure
                    self._add_html_formatted_text_to_paragraph(content_p, para)
                    
                    # Match CSS spacing
                    content_p.space_after = Pt(8)  # Slightly less than default for tighter layout
        
        elif content_lists:
            # Process lists with exact HTML structure
            for list_elem in content_lists:
                # Add list items as properly formatted paragraphs
                items = list_elem.find_all('li')
                for item in items:
                    item_p = cell.add_paragraph()
                    self._add_html_formatted_text_to_paragraph(item_p, item)
                    item_p.space_after = Pt(4)
        
        else:
            # Fallback: single text block
            content_p = cell.paragraphs[0]
            content_run = content_p.add_run(content_text)
            content_run.font.size = Pt(18)  # Matches CSS font-size: 18px
            content_run.font.name = self.primary_font
            content_run.font.color.rgb = self.robeco_colors['text_dark']  # Matches CSS color: var(--text-dark)
    
    def _add_html_formatted_text_to_paragraph(self, paragraph, html_element):
        """Add text with EXACT HTML formatting preservation (bold, italic, etc.)"""
        logger.info(f"🎨 HTML FORMATTING: Processing {html_element.name} with precise style replication")
        
        # Process all content including nested tags
        for element in html_element.contents:
            if hasattr(element, 'name') and element.name:
                element_text = element.get_text()
                
                if element.name == 'strong' or element.name == 'b':
                    # Bold text - matches HTML <strong> or <b>
                    run = paragraph.add_run(element_text)
                    run.font.bold = True
                    run.font.size = Pt(18)
                    run.font.name = self.primary_font
                    run.font.color.rgb = self.robeco_colors['text_dark']
                elif element.name == 'em' or element.name == 'i':
                    # Italic text - matches HTML <em> or <i>
                    run = paragraph.add_run(element_text)
                    run.font.italic = True
                    run.font.size = Pt(18)
                    run.font.name = self.primary_font
                    run.font.color.rgb = self.robeco_colors['text_dark']
                else:
                    # Regular text
                    run = paragraph.add_run(element_text)
                    run.font.size = Pt(18)
                    run.font.name = self.primary_font
                    run.font.color.rgb = self.robeco_colors['text_dark']
            else:
                # Plain text node
                text = str(element).strip()
                if text:
                    run = paragraph.add_run(text)
                    run.font.size = Pt(18)  # Matches CSS .content-item font-size: 18px
                    run.font.name = self.primary_font
                    run.font.color.rgb = self.robeco_colors['text_dark']  # Matches CSS color: var(--text-dark)

    def _process_bullet_list_square(self, doc: Document, list_soup):
        """Process bullet-list-square with special formatting"""
        logger.info("🔍 Processing BULLET-LIST-SQUARE")
        
        # Look for h4 headers within the list
        headers = list_soup.find_all('h4')
        for header in headers:
            p = doc.add_paragraph()
            run = p.add_run(header.get_text().strip())
            run.font.size = Pt(20)  # Matches 22.5px from CSS
            run.font.name = self.primary_font
            run.font.bold = True
            run.font.color.rgb = self.robeco_colors['blue_darker']
            p.space_after = Pt(6)
        
        # Process paragraphs
        paragraphs = list_soup.find_all('p')
        for para in paragraphs:
            p = doc.add_paragraph()
            self._add_formatted_text_to_paragraph(p, para)
            # Use intelligent alignment for bullet list paragraphs
            self._fix_element_alignment(p, para)
            p.space_after = Pt(8)
        
        # Process ul/ol lists within
        lists = list_soup.find_all(['ul', 'ol'])
        for list_elem in lists:
            self._add_list(doc, list_elem)

    def _add_formatted_text_to_paragraph(self, paragraph, html_element):
        """Add formatted text from HTML element to Word paragraph, preserving bold/italic"""
        
        logger.info(f"🔍 Processing formatted text: {html_element.name} with {len(html_element.contents)} children")
        total_text_added = ""
        
        # Process all content including nested formatting
        for i, element in enumerate(html_element.contents):
            if hasattr(element, 'name') and element.name:
                element_text = element.get_text()
                logger.info(f"📝 Processing element {i}: {element.name} = '{element_text[:50]}...'")
                
                if element.name == 'strong':
                    run = paragraph.add_run(element_text)
                    run.font.bold = True
                    run.font.size = Pt(16)
                    run.font.name = self.primary_font
                    total_text_added += element_text
                elif element.name == 'em':
                    run = paragraph.add_run(element_text)
                    run.font.italic = True
                    run.font.size = Pt(16)
                    run.font.name = self.primary_font
                    total_text_added += element_text
                else:
                    # Other tags, get text content
                    run = paragraph.add_run(element_text)
                    run.font.size = Pt(16)
                    run.font.name = self.primary_font
                    total_text_added += element_text
            else:
                # Plain text node
                text = str(element).strip()
                if text:
                    logger.info(f"📝 Processing text node {i}: '{text[:50]}...'")
                    run = paragraph.add_run(text)
                    run.font.size = Pt(16)
                    run.font.name = self.primary_font
                    total_text_added += text
        
        logger.info(f"✅ Added {len(total_text_added)} characters to paragraph")
    
    # DEPRECATED: Old method replaced by _force_precise_flexbox_layout
    
    def _add_analysis_item_border(self, table):
        """Add bottom border to analysis item table to match HTML border-bottom: 3.5px solid #E0E0E0"""
        try:
            # Get table properties
            tbl_pr = table._element.find(qn('w:tblPr'))
            if tbl_pr is None:
                tbl_pr = OxmlElement('w:tblPr')
                table._element.insert(0, tbl_pr)
            
            # Add table borders
            tbl_borders = tbl_pr.find(qn('w:tblBorders'))
            if tbl_borders is None:
                tbl_borders = OxmlElement('w:tblBorders')
                tbl_pr.append(tbl_borders)
            
            # Add bottom border (matches CSS border-bottom: 3.5px solid #E0E0E0)
            bottom_border = OxmlElement('w:bottom')
            bottom_border.set(qn('w:val'), 'single')
            bottom_border.set(qn('w:sz'), '21')  # 3.5px equivalent (21 = 3.5px * 6)
            bottom_border.set(qn('w:color'), 'E0E0E0')  # Light gray
            tbl_borders.append(bottom_border)
            
            # Add top border for first analysis item (matches CSS border-top: 5px solid var(--robeco-blue))
            if hasattr(self, '_first_analysis_item') and not self._first_analysis_item:
                top_border = OxmlElement('w:top')
                top_border.set(qn('w:val'), 'single')
                top_border.set(qn('w:sz'), '30')  # 5px equivalent
                top_border.set(qn('w:color'), '005F90')  # Robeco blue
                tbl_borders.append(top_border)
                self._first_analysis_item = True
                logger.info("🔵 Added blue top border for first analysis item")
            
            logger.info("📏 Added analysis item borders to match HTML styling")
            
        except Exception as e:
            logger.error(f"❌ Error adding analysis item border: {e}")
    
    def _add_orange_separator_line(self, doc: Document):
        """Add orange separator line to match HTML .orange-separator"""
        try:
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Add orange top border to paragraph
            p_pr = p._element.get_or_add_pPr()
            p_bdr = OxmlElement('w:pBdr')
            top_border = OxmlElement('w:top')
            top_border.set(qn('w:val'), 'single')
            top_border.set(qn('w:sz'), '42')  # 7px equivalent (7 * 6)
            top_border.set(qn('w:color'), 'FF8C00')  # Orange color
            p_bdr.append(top_border)
            p_pr.append(p_bdr)
            
            p.space_after = Pt(12)
            logger.info("🟠 Added orange separator line")
            
        except Exception as e:
            logger.error(f"❌ Error adding orange separator: {e}")
    
    def _process_all_images_in_slide(self, doc: Document, slide_soup):
        """Process all images in slide with ENHANCED positioning and context analysis"""
        try:
            all_images = slide_soup.find_all('img')
            logger.info(f"📸 ENHANCED: Found {len(all_images)} images in slide")
            
            # Also look for chart containers with canvas/svg elements or chart data
            chart_containers = slide_soup.find_all(['canvas', 'svg'])
            chart_divs = slide_soup.find_all('div', class_=lambda x: x and ('chart' in ' '.join(x) or 'graph' in ' '.join(x) or 'visualization' in ' '.join(x)))
            chart_containers.extend(chart_divs)
            
            if chart_containers:
                logger.info(f"📊 Found {len(chart_containers)} chart containers (canvas/svg/chart-divs)")
                for container in chart_containers:
                    self._process_chart_container(doc, container, slide_soup)
            else:
                logger.info("📊 No chart containers found in this slide")
            
            # 📍 CREATE IMAGE POSITION MAP: Analyze HTML structure for precise positioning
            image_position_map = self._create_image_position_map(slide_soup, all_images)
            
            for i, img in enumerate(all_images):
                img_src = img.get('src', '')
                img_alt = img.get('alt', '')
                img_classes = img.get('class', [])
                
                logger.info(f"📸 Processing image {i+1}/{len(all_images)}: '{img_alt}' at {img_src}")
                
                # Skip placeholders and empty sources
                if not img_src or 'placehold.co' in img_src or '[COMPANY_LOGO]' in img_src:
                    logger.info("🔄 Skipping placeholder/template image")
                    continue
                
                # 🏢 CATEGORIZE IMAGE TYPE for proper positioning
                image_type = self._determine_image_type(img, img_src, img_alt, img_classes)
                
                if image_type == 'robeco_logo':
                    logger.info("🏢 Robeco logo - inserting into header section")
                    self._insert_logo_image(doc, img, img_src, img_alt, 'robeco_logo')
                elif image_type == 'company_icon':
                    logger.info("🏢 Company icon - inserting into company header section")
                    self._insert_logo_image(doc, img, img_src, img_alt, 'company_icon')
                elif image_type == 'chart':
                    logger.info("📊 Chart image - inserting with chart-specific positioning")
                    self._insert_chart_image_with_context(doc, img, img_src, img_alt, image_position_map.get(i, {}))
                elif image_type == 'content_image':
                    logger.info("📷 Content image - inserting with content-specific positioning")
                    self._insert_content_image_with_context(doc, img, img_src, img_alt, image_position_map.get(i, {}))
                else:
                    logger.info("🖼️ Generic image - using contextual insertion")
                    self._insert_contextual_image(doc, img, img_src, img_alt)
                    
        except Exception as e:
            logger.error(f"❌ Enhanced image processing failed: {e}")
    
    def _create_image_position_map(self, slide_soup, images):
        """Create a detailed position map for images based on HTML structure"""
        position_map = {}
        
        for i, img in enumerate(images):
            try:
                # Analyze parent hierarchy
                parent = img.parent
                parent_classes = parent.get('class', []) if parent else []
                
                # Find containing section
                section_container = img
                while section_container and section_container.parent:
                    section_container = section_container.parent
                    section_classes = section_container.get('class', []) if hasattr(section_container, 'get') else []
                    if any(cls in section_classes for cls in ['slide', 'content-block', 'analysis-item', 'metrics-grid']):
                        break
                
                position_info = {
                    'parent_classes': parent_classes,
                    'section_classes': section_classes if 'section_classes' in locals() else [],
                    'position_in_slide': i,
                    'total_images': len(images)
                }
                
                position_map[i] = position_info
                logger.info(f"📍 Image {i} position: parent={parent_classes}, section={position_info.get('section_classes', [])}")
                
            except Exception as e:
                logger.warning(f"⚠️ Could not map position for image {i}: {e}")
                position_map[i] = {}
        
        return position_map
    
    def _determine_image_type(self, img, img_src: str, img_alt: str, img_classes: list) -> str:
        """Determine image type for precise positioning"""
        # Check for Robeco branding
        if 'robeco' in img_src.lower() or 'robeco' in img_alt.lower():
            return 'robeco_logo'
        
        # Check for company icons (Clearbit, etc.)
        if ('clearbit.com' in img_src or 'icon' in img_classes or 
            'icon' in img_alt.lower() or 'company' in img_alt.lower()):
            return 'company_icon'
        
        # Check for charts and data visualizations
        if ('chart' in img_classes or 'chart' in img_alt.lower() or 
            'graph' in img_alt.lower() or 'visualization' in img_alt.lower()):
            return 'chart'
        
        # Check parent context for chart containers
        parent = img.parent
        if parent:
            parent_classes = parent.get('class', [])
            if any(cls in parent_classes for cls in ['chart-container', 'stock-chart-container', 'visualization']):
                return 'chart'
        
        # Default to content image
        return 'content_image'
    
    def _insert_chart_image_with_context(self, doc: Document, img_element, img_src: str, img_alt: str, position_info: dict):
        """Insert chart image with precise chart-specific positioning"""
        try:
            # Charts typically center-aligned and larger
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Larger size for charts
            target_width = Inches(5.0)  # Charts need to be readable
            
            success = self._download_and_insert_image_inline(p, img_src, img_alt, target_width)
            
            if success:
                # Add chart caption if meaningful
                if img_alt and len(img_alt) > 3 and 'chart' not in img_alt.lower():
                    caption_p = doc.add_paragraph(f"Figure: {img_alt}")
                    caption_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    caption_run = caption_p.runs[0]
                    caption_run.font.size = Pt(10)
                    caption_run.font.italic = True
                    caption_run.font.color.rgb = self.robeco_colors['text_secondary']
                
                logger.info(f"✅ Chart image inserted: {img_alt}")
            else:
                # Remove failed paragraph
                doc.paragraphs.pop()
                
        except Exception as e:
            logger.error(f"❌ Chart image insertion failed: {e}")
    
    def _insert_content_image_with_context(self, doc: Document, img_element, img_src: str, img_alt: str, position_info: dict):
        """Insert content image with context-aware positioning"""
        try:
            # Determine alignment based on context
            parent_classes = position_info.get('parent_classes', [])
            
            p = doc.add_paragraph()
            
            if any(cls in parent_classes for cls in ['center', 'text-center']):
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                target_width = Inches(3.0)
            elif any(cls in parent_classes for cls in ['right', 'text-right']):
                p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                target_width = Inches(2.5)
            else:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                target_width = Inches(2.0)
            
            success = self._download_and_insert_image_inline(p, img_src, img_alt, target_width)
            
            if success:
                logger.info(f"✅ Content image inserted: {img_alt}")
            else:
                doc.paragraphs.pop()
                
        except Exception as e:
            logger.error(f"❌ Content image insertion failed: {e}")
    
    def _generate_intelligent_logo_fallbacks(self, original_url: str, company_name: str) -> list:
        """Generate intelligent logo fallback URLs for any company"""
        fallback_urls = [original_url]  # Always try original first
        
        try:
            # Extract company name parts for intelligent URL generation
            company_clean = company_name.lower().strip()
            
            # Remove common suffixes and words
            suffixes_to_remove = [
                'inc', 'inc.', 'corp', 'corp.', 'corporation', 'company', 'co', 'co.', 
                'ltd', 'ltd.', 'limited', 'plc', 'llc', 'trust', 'reit', 'group', 
                'holdings', 'holding', 'international', 'intl', 'global', 'worldwide',
                'logistics', 'commercial', 'property', 'properties', 'real estate',
                '&', 'and', 'the'
            ]
            
            words = company_clean.replace('&', '').replace(',', '').split()
            filtered_words = []
            
            for word in words:
                word_clean = word.strip('.,()[]{}')
                if word_clean and word_clean not in suffixes_to_remove:
                    filtered_words.append(word_clean)
            
            # Generate various domain combinations
            if len(filtered_words) >= 1:
                # Single word domains
                main_word = filtered_words[0]
                fallback_urls.extend([
                    f"https://logo.clearbit.com/{main_word}.com",
                    f"https://logo.clearbit.com/{main_word}.co.uk",
                    f"https://logo.clearbit.com/{main_word}.sg",
                    f"https://logo.clearbit.com/{main_word}.com.au"
                ])
                
                # Two word combinations
                if len(filtered_words) >= 2:
                    two_words = filtered_words[0] + filtered_words[1]
                    fallback_urls.extend([
                        f"https://logo.clearbit.com/{two_words}.com",
                        f"https://logo.clearbit.com/{filtered_words[0]}{filtered_words[1]}.com",
                        f"https://logo.clearbit.com/{filtered_words[0]}-{filtered_words[1]}.com",
                        f"https://logo.clearbit.com/{filtered_words[0]}_{filtered_words[1]}.com"
                    ])
                
                # Three word combinations (for companies like "Frasers Logistics Commercial")
                if len(filtered_words) >= 3:
                    three_words = filtered_words[0] + filtered_words[1] + filtered_words[2]
                    fallback_urls.extend([
                        f"https://logo.clearbit.com/{three_words}.com",
                        f"https://logo.clearbit.com/{filtered_words[0]}{filtered_words[1]}{filtered_words[2]}.com"
                    ])
            
            # Stock ticker-based fallbacks (extract from URL or use common patterns)
            if 'clearbit.com/' in original_url:
                original_domain = original_url.split('clearbit.com/')[-1]
                if original_domain and '.' not in original_domain:
                    # Try common extensions for the domain
                    for ext in ['.com', '.co.uk', '.sg', '.com.au', '.de', '.fr', '.nl']:
                        fallback_urls.append(f"https://logo.clearbit.com/{original_domain}{ext}")
            
            # Remove duplicates while preserving order
            seen = set()
            unique_fallbacks = []
            for url in fallback_urls:
                if url not in seen:
                    seen.add(url)
                    unique_fallbacks.append(url)
            
            logger.info(f"🎯 Generated {len(unique_fallbacks)} intelligent logo fallbacks for '{company_name}'")
            return unique_fallbacks[:10]  # Limit to 10 attempts
            
        except Exception as e:
            logger.warning(f"⚠️ Error generating logo fallbacks: {e}")
            return [original_url]  # Return at least the original URL
    
    def _insert_contextual_image(self, doc: Document, img_element, img_src: str, img_alt: str):
        """Insert image at its contextual position based on HTML structure"""
        try:
            # Determine image context based on parent elements
            parent = img_element.parent
            parent_classes = parent.get('class', []) if parent else []
            
            logger.info(f"📸 Image context - parent: {parent.name if parent else 'None'}, classes: {parent_classes}")
            
            # Create paragraph for the image
            p = doc.add_paragraph()
            
            # Determine alignment based on context
            if 'chart' in parent_classes or 'chart' in img_alt.lower():
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                target_width = Inches(4.0)  # Charts are typically larger
                logger.info("📈 Chart image - center aligned, large size")
            elif 'icon' in parent_classes or 'logo' in parent_classes:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                target_width = Inches(0.5)  # Icons are small
                logger.info("📷 Icon image - left aligned, small size")
            else:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                target_width = Inches(2.0)  # Default size
                logger.info("📷 Generic image - center aligned, medium size")
            
            # Try to download and insert the image
            success = self._download_and_insert_image_inline(p, img_src, img_alt, target_width)
            
            if success:
                logger.info(f"✅ Successfully inserted contextual image: {img_alt}")
                # Add caption if alt text is meaningful
                if img_alt and len(img_alt) > 3 and 'icon' not in img_alt.lower():
                    caption_p = doc.add_paragraph(img_alt)
                    caption_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    caption_run = caption_p.runs[0]
                    caption_run.font.size = Pt(10)
                    caption_run.font.italic = True
                    caption_run.font.color.rgb = self.robeco_colors['text_secondary']
            else:
                # Remove empty paragraph if image failed
                doc.paragraphs.pop()
                logger.warning(f"⚠️ Failed to insert image, paragraph removed: {img_src}")
                
        except Exception as e:
            logger.error(f"❌ Error inserting contextual image: {e}")
    
    # Removed _add_robeco_logo method - functionality moved to _add_report_header
    
    def _add_company_icon(self, doc: Document, header_soup) -> bool:
        """Add company icon from HTML - positioned next to company name"""
        try:
            # Look for company icon (Clearbit logo)
            img_tags = header_soup.find_all('img')
            
            for img in img_tags:
                img_src = img.get('src', '')
                img_alt = img.get('alt', '')
                img_classes = img.get('class', [])
                
                logger.info(f"📸 Found img: src='{img_src}', alt='{img_alt}', classes='{img_classes}'")
                
                # Check if it's a company icon (clearbit or has 'icon' class)
                if ('clearbit.com' in img_src or 'icon' in img_classes or 
                    'icon' in img_alt.lower() or img_src.endswith('.co.jp')):
                    
                    # Create inline company icon with name
                    company_name_elem = header_soup.find('h1', class_='name')
                    if company_name_elem:
                        p = doc.add_paragraph()
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        
                        # Try intelligent fallback URLs for company logos
                        success = False
                        company_name = company_name_elem.get_text().strip()
                        fallback_urls = self._generate_intelligent_logo_fallbacks(img_src, company_name)
                        
                        for i, url in enumerate(fallback_urls):
                            if i > 0:  # Don't log for original URL
                                logger.info(f"🔄 Trying intelligent fallback {i}: {url}")
                            success = self._download_and_insert_image_inline(p, url, img_alt, Inches(0.4))
                            if success:
                                logger.info(f"✅ Success with fallback URL {i}: {url}")
                                break
                        
                        # Add company name after icon (whether icon succeeded or not)
                        run = p.add_run((" " if success else "") + company_name_elem.get_text().strip())
                        run.font.size = Pt(32)
                        run.font.name = self.primary_font
                        run.font.bold = True
                        run.font.color.rgb = self.robeco_colors['text_dark']
                        p.space_after = Pt(8)
                        
                        if success:
                            logger.info(f"✅ Added company icon with name: {img_alt}")
                        else:
                            logger.info(f"📝 Added company name without icon (logo failed): {company_name_elem.get_text().strip()}")
                        return True
                    else:
                        # No company name found, try to add icon only
                        p = doc.add_paragraph()
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        success = self._download_and_insert_image_inline(p, img_src, img_alt, Inches(0.4))
                        if success:
                            logger.info(f"✅ Added company icon: {img_alt}")
                            return True
                        else:
                            # Remove empty paragraph if no icon and no name
                            doc.paragraphs.pop()
            
            logger.info("ℹ️ No company icon found")
            return False
            
        except Exception as e:
            logger.error(f"❌ Error adding company icon: {e}")
            return False
    
    def _insert_logo_image(self, doc: Document, img_element, img_src: str, img_alt: str, logo_type: str):
        """Insert logo images (Robeco logo or company icon) into Word document"""
        try:
            logger.info(f"🖼️ Inserting {logo_type}: {img_alt}")
            
            # Create paragraph for the logo
            p = doc.add_paragraph()
            
            # Configure alignment and size based on logo type
            if logo_type == 'robeco_logo':
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                target_width = Inches(1.2)  # Robeco logo size
                logger.info("🏢 Robeco logo - left aligned")
            elif logo_type == 'company_icon':
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  
                target_width = Inches(0.4)  # Company icon size
                logger.info("🏢 Company icon - center aligned")
            else:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                target_width = Inches(0.5)  # Default size
            
            # Download and insert the image
            success = self._download_and_insert_image_inline(p, img_src, img_alt, target_width)
            
            if success:
                logger.info(f"✅ Successfully inserted {logo_type}: {img_alt}")
            else:
                # Remove empty paragraph if image failed
                doc.paragraphs.pop()
                logger.warning(f"⚠️ Failed to insert {logo_type}, paragraph removed: {img_src}")
                
        except Exception as e:
            logger.error(f"❌ Error inserting {logo_type}: {e}")
    
    def _process_chart_container(self, doc: Document, container, slide_soup):
        """Process chart containers (canvas/svg/chart divs) and extract chart data"""
        try:
            logger.info(f"📊 Processing chart container: {container.name}, classes: {container.get('class', [])}")
            
            # Look for stock data in nearby script tags
            script_tags = slide_soup.find_all('script')
            chart_data = None
            
            for script in script_tags:
                script_text = script.get_text()
                if 'stockData' in script_text or 'chartData' in script_text:
                    chart_data = self._extract_chart_data_from_script(script_text)
                    break
            
            if chart_data:
                logger.info(f"📈 Extracted chart data: {len(chart_data.get('prices', []))} price points")
                self._insert_chart_summary(doc, chart_data)
            else:
                # Fallback: Create placeholder for chart
                logger.info("📊 No chart data found, creating chart placeholder")
                self._insert_chart_placeholder(doc, container)
                
        except Exception as e:
            logger.error(f"❌ Error processing chart container: {e}")
    
    def _extract_chart_data_from_script(self, script_text: str) -> dict:
        """Extract chart data from JavaScript code"""
        try:
            import re
            
            # Extract price data
            price_matches = re.findall(r"'price':\s*([\d.]+)", script_text)
            date_matches = re.findall(r"'date':\s*'([\d-]+)'", script_text)
            
            # Extract ticker or company name
            ticker_match = re.search(r"'ticker':\s*'([^']+)'", script_text)
            company_match = re.search(r"'company':\s*'([^']+)'", script_text)
            
            if price_matches and date_matches:
                prices = [float(p) for p in price_matches]
                current_price = prices[-1]
                start_price = prices[0]
                price_change = current_price - start_price
                price_change_pct = (price_change / start_price) * 100
                
                return {
                    'prices': prices,
                    'dates': date_matches,
                    'current_price': current_price,
                    'start_price': start_price,
                    'price_change': price_change,
                    'price_change_pct': price_change_pct,
                    'ticker': ticker_match.group(1) if ticker_match else 'Unknown',
                    'company': company_match.group(1) if company_match else 'Company'
                }
            
            return {}
            
        except Exception as e:
            logger.error(f"❌ Error extracting chart data: {e}")
            return {}
    
    def _insert_chart_summary(self, doc: Document, chart_data: dict):
        """Insert a text-based chart summary"""
        try:
            # Chart title
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title_run = p.add_run(f"📈 {chart_data.get('ticker', 'Stock')} Price Chart")
            title_run.font.size = Pt(14)
            title_run.font.bold = True
            title_run.font.color.rgb = self.robeco_colors['brown_black']
            
            # Price summary
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            current_price = chart_data.get('current_price', 0)
            price_change = chart_data.get('price_change', 0)
            price_change_pct = chart_data.get('price_change_pct', 0)
            
            # Color based on performance
            color = self.robeco_colors['success'] if price_change >= 0 else self.robeco_colors['error']
            arrow = "▲" if price_change >= 0 else "▼"
            
            summary_text = f"Current: ${current_price:.2f} {arrow} ${abs(price_change):.2f} ({price_change_pct:+.1f}%)"
            summary_run = p.add_run(summary_text)
            summary_run.font.size = Pt(12)
            summary_run.font.color.rgb = color
            
            logger.info(f"✅ Inserted chart summary for {chart_data.get('ticker', 'Unknown')}")
            
        except Exception as e:
            logger.error(f"❌ Error inserting chart summary: {e}")
    
    def _insert_chart_placeholder(self, doc: Document, container):
        """Insert a placeholder for missing charts"""
        try:
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            placeholder_run = p.add_run("📊 [Chart - Data visualization]")
            placeholder_run.font.size = Pt(12)
            placeholder_run.font.italic = True
            placeholder_run.font.color.rgb = self.robeco_colors['text_secondary']
            
            logger.info("✅ Inserted chart placeholder")
            
        except Exception as e:
            logger.error(f"❌ Error inserting chart placeholder: {e}")
    
    def _add_full_width_table(self, doc: Document, table_soup):
        """Add table with full page width (for prose layout)"""
        try:
            logger.info("📊 Creating full-width table")
            rows = table_soup.find_all('tr')
            if not rows:
                return
            
            # Create Word table with full width
            max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
            table = doc.add_table(rows=len(rows), cols=max_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Set table to use full page width
            table.autofit = False
            col_width_each = Inches(7.5 / max_cols) if max_cols > 0 else Inches(1.5)
            for i, col in enumerate(table.columns):
                col.width = col_width_each  # Distribute full width evenly
            
            # Fill table content
            for i, row_soup in enumerate(rows):
                cells = row_soup.find_all(['td', 'th'])
                for j, cell_soup in enumerate(cells):
                    if j < max_cols:
                        cell = table.cell(i, j)
                        cell_p = cell.paragraphs[0]
                        cell_p.clear()
                        
                        # Add cell content with proper formatting
                        cell_text = cell_soup.get_text().strip()
                        run = cell_p.add_run(cell_text)
                        
                        # Header styling for th elements
                        if cell_soup.name == 'th':
                            run.font.bold = True
                            run.font.color.rgb = self.robeco_colors['blue_darker']
                            cell_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        else:
                            cell_p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        
                        run.font.size = Pt(12)
                        run.font.name = self.primary_font
                        
                        # Add cell padding
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            
            # Add spacing after table
            spacing_para = doc.add_paragraph()
            spacing_para.space_after = Pt(16)
            
            logger.info(f"✅ Added full-width table with {len(rows)} rows, {max_cols} columns")
            
        except Exception as e:
            logger.error(f"❌ Error adding full-width table: {e}")
    
    def _add_full_width_intro_content(self, doc: Document, container_soup):
        """Process intro-chart container as full-width content (for prose layout)"""
        try:
            logger.info("📝 Processing intro-chart container as full-width content")
            
            # Extract intro text block
            intro_block = container_soup.find(class_='intro-text-block')
            if intro_block:
                logger.info("✅ Found intro-text-block, processing as full-width")
                for para in intro_block.find_all('p'):
                    self._add_paragraph(doc, para)
            
            # Extract any other content 
            for child in container_soup.children:
                if hasattr(child, 'name') and child.name is not None:
                    if child.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                        self._add_header(doc, child)
                    elif child.name == 'p':
                        self._add_paragraph(doc, child)
                    elif child.name == 'table':
                        self._add_full_width_table(doc, child)
            
            logger.info("✅ Completed full-width intro content processing")
            
        except Exception as e:
            logger.error(f"❌ Error processing full-width intro content: {e}")
    
    def _process_slide_logo(self, doc: Document, logo_container):
        """Process slide-logo container with Robeco logo"""
        try:
            logger.info("🏢 Processing slide-logo container")
            
            # Find Robeco logo image
            img = logo_container.find('img')
            if img:
                img_src = img.get('src', '')
                img_alt = img.get('alt', '')
                
                if 'robeco' in img_src.lower() or 'robeco' in img_alt.lower():
                    logger.info("🏢 Found Robeco logo in slide-logo container")
                    
                    # Create paragraph for logo
                    p = doc.add_paragraph()
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    
                    # Insert Robeco logo with appropriate size
                    success = self._download_and_insert_image_inline(p, img_src, img_alt, Inches(1.2))
                    
                    if success:
                        logger.info("✅ Successfully inserted Robeco logo from slide-logo container")
                    else:
                        # Remove empty paragraph if failed
                        doc.paragraphs.pop()
                        logger.warning("⚠️ Failed to insert Robeco logo from slide-logo container")
                else:
                    logger.info("ℹ️ Image in slide-logo container is not Robeco logo")
            else:
                logger.info("ℹ️ No image found in slide-logo container")
                
        except Exception as e:
            logger.error(f"❌ Error processing slide-logo container: {e}")
    
    def _add_slide_logo_header(self, doc: Document, slide_logo_soup):
        """Add slide logo at top-right following HTML structure"""
        try:
            logger.info("🏢 Adding slide logo header (top-right positioning)")
            
            # Find Robeco logo image in slide-logo container
            img = slide_logo_soup.find('img')
            if img:
                img_src = img.get('src', '')
                img_alt = img.get('alt', '')
                
                if 'robeco' in img_src.lower() or 'robeco' in img_alt.lower():
                    # Create right-aligned paragraph for top-right positioning
                    p = doc.add_paragraph()
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    
                    # Insert Robeco logo with appropriate size (2rem = ~24px = ~0.33 inches)
                    success = self._download_and_insert_image_inline(p, img_src, img_alt, Inches(1.0))
                    
                    if success:
                        logger.info("✅ Successfully added slide logo header (top-right)")
                        # Add spacing after logo
                        p.space_after = Pt(12)
                    else:
                        # Remove empty paragraph if failed
                        doc.paragraphs.pop()
                        logger.warning("⚠️ Failed to add slide logo header")
                else:
                    logger.info("ℹ️ Image in slide-logo is not Robeco logo")
            else:
                logger.info("ℹ️ No image found in slide-logo container")
                
        except Exception as e:
            logger.error(f"❌ Error adding slide logo header: {e}")
    
    def _download_and_insert_image_inline(self, paragraph, img_url: str, alt_text: str, target_width: object) -> bool:
        """Download image from URL and insert into Word document"""
        try:
            # Check cache first
            if img_url in self._image_cache:
                logger.info(f"📸 Using cached image: {img_url}")
                image_data = self._image_cache[img_url]
            else:
                # Download image
                logger.info(f"📸 Downloading image: {img_url}")
                response = requests.get(img_url, timeout=10, headers={
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
                })
                response.raise_for_status()
                
                image_data = response.content
                self._image_cache[img_url] = image_data
                logger.info(f"✅ Downloaded {len(image_data)} bytes")
            
            # Insert image inline into the provided paragraph
            image_stream = BytesIO(image_data)
            
            # Try to determine optimal size
            try:
                with Image.open(image_stream) as pil_img:
                    width, height = pil_img.size
                    # Use target width, maintain aspect ratio
                    if width > 0:
                        aspect_ratio = height / width
                        img_width = target_width
                        img_height = Inches(target_width.inches * aspect_ratio)
                    else:
                        img_width = target_width
                        img_height = target_width
                        
                    logger.info(f"📐 Image dimensions: {width}x{height} -> {img_width.inches:.2f}x{img_height.inches:.2f} inches")
            except Exception as e:
                logger.warning(f"⚠️ Could not determine image size: {e}, using defaults")
                img_width = target_width
                img_height = target_width
            
            # Reset stream position
            image_stream.seek(0)
            
            # Add the image inline
            run = paragraph.add_run()
            run.add_picture(image_stream, width=img_width, height=img_height)
            
            logger.info(f"✅ Successfully inserted image: {alt_text}")
            return True
            
        except Exception as e:
            logger.error(f"❌ Failed to download/insert image from {img_url}: {e}")
            return False
    
    def _fix_element_alignment(self, paragraph, html_element):
        """Fix element alignment to match HTML positioning with ENHANCED CSS analysis"""
        try:
            # 🔍 DEEP HTML STRUCTURE ANALYSIS
            element_classes = html_element.get('class', [])
            parent = html_element.parent
            parent_classes = parent.get('class', []) if parent else []
            
            # Check grandparent for layout context
            grandparent_classes = []
            if parent and hasattr(parent, 'parent') and parent.parent:
                grandparent_classes = parent.parent.get('class', []) if hasattr(parent.parent, 'get') else []
            
            # Default alignment based on HTML context
            alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            # 🎯 ADVANCED ALIGNMENT DETECTION
            
            # Check for explicit center alignment
            center_indicators = ['center', 'text-center', 'align-center', 'company-header', 'report-header', 'report-title', 'report-subtitle']
            if any(cls in element_classes + parent_classes + grandparent_classes for cls in center_indicators):
                alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                logger.info(f"🎯 CENTER alignment: {element_classes}")
            
            # Check for right alignment
            elif any(cls in element_classes + parent_classes for cls in ['right', 'text-right', 'align-right']):
                alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                logger.info(f"🎯 RIGHT alignment: {element_classes}")
            
            # 📄 CONTENT-SPECIFIC ALIGNMENT RULES
            
            # Content items in analysis tables should justify for better readability
            elif 'content-item' in parent_classes or 'intro-text-block' in element_classes:
                alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                logger.info(f"🎯 JUSTIFY alignment for content: {element_classes}")
            
            # Titles and headers typically left-align in analysis items
            elif 'item-title' in element_classes or html_element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                logger.info(f"🎯 LEFT alignment for title/header: {element_classes}")
            
            # 🏢 LAYOUT CONTEXT DETECTION
            
            # Elements in flexbox layouts (analysis-item) maintain left/justify alignment
            elif 'analysis-item' in parent_classes or 'analysis-item' in grandparent_classes:
                # For analysis items, content should justify, titles should left-align
                if 'content-item' in element_classes or html_element.name == 'p':
                    alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                else:
                    alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                logger.info(f"🎯 Analysis item context alignment: {alignment}")
            
            # Company and report headers center
            elif any(cls in parent_classes + grandparent_classes for cls in ['company-header', 'report-header-container', 'slide-header']):
                alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                logger.info(f"🎯 Header context CENTER: {parent_classes}")
            
            # Default left alignment
            else:
                alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                logger.info(f"🎯 DEFAULT LEFT alignment: {element_classes}")
            
            paragraph.alignment = alignment
            return alignment
            
        except Exception as e:
            logger.warning(f"⚠️ Alignment detection failed: {e}, using LEFT")
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            return WD_PARAGRAPH_ALIGNMENT.LEFT

# Global instance for use across the application
word_report_generator = RobecoWordReportGenerator()